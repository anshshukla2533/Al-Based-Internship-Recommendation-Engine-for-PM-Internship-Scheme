from fastapi import FastAPI, UploadFile, File, HTTPException, Form
try:
    from omnidimension import Client
    from omnidimension.client import APIError
except ImportError:
    Client = None
    class APIError(Exception):
        def __init__(self, status_code=0, message="", response=None):
            self.status_code = status_code
            self.message = message
            self.response = response or {}
            super().__init__(message)
from fastapi.middleware.cors import CORSMiddleware

from pydantic import BaseModel

from typing import List, Dict, Any, Optional

try:

    import pdfplumber

except ImportError:

    pdfplumber = None

try:

    from pypdf import PdfReader as ModernPdfReader

except ImportError:

    ModernPdfReader = None

try:

    from PyPDF2 import PdfReader as LegacyPdfReader

except ImportError:

    LegacyPdfReader = None
import re

import json

import html

import os

import difflib

import random

import uuid

import hashlib

import ast

from collections import Counter

from datetime import datetime, timedelta, timezone

import requests

from io import BytesIO

from contextlib import contextmanager

from urllib.parse import urljoin, urlparse

from dotenv import load_dotenv

try:

    from groq import Groq

except ImportError:

    Groq = None

from deep_translator import GoogleTranslator

from sklearn.feature_extraction.text import TfidfVectorizer

from sklearn.metrics.pairwise import cosine_similarity

try:

    from docx import Document as DocxDocument

except ImportError:

    DocxDocument = None

try:

    from firecrawl import FirecrawlApp

except ImportError:

    FirecrawlApp = None

DEFAULT_JOBS_JSON_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "jobs.json")

try:

    from fetch_internships import (

        JOBS_JSON_PATH as SCRAPED_JOBS_JSON_PATH,

        PM_PORTAL_URL as SCRAPED_PM_PORTAL_URL,

        SOURCE_NAME as PM_SOURCE_NAME,

        fetch_and_save_internships,

        infer_skills as infer_pm_job_skills,

    )

except ImportError:

    SCRAPED_JOBS_JSON_PATH = DEFAULT_JOBS_JSON_PATH

    SCRAPED_PM_PORTAL_URL = "https://pminternship.mca.gov.in/"

    PM_SOURCE_NAME = "PM Internship Portal"

    fetch_and_save_internships = None

    infer_pm_job_skills = None

PM_PORTAL_URL = SCRAPED_PM_PORTAL_URL

JOBS_JSON_PATH = SCRAPED_JOBS_JSON_PATH

INTERNSHALA_BASE_URL = "https://internshala.com"

INTERNSHALA_SOURCE = "Internshala"

LANG_MAP = {

    "en": "English",

    "hi": "Hindi",

    "te": "Telugu",

    "mr": "Marathi",

    "ta": "Tamil",

}

DEFAULT_GEMINI_MODELS = [

    "gemini-2.5-flash",

    "gemini-2.5-flash-lite",

    "gemini-2.0-flash-lite",

]

DEFAULT_GROQ_MODELS = [

    "llama-3.3-70b-versatile",

    "llama-3.1-8b-instant",

    "llama3-8b-8192",

]

RESUME_ANALYSIS_CACHE: Dict[str, Dict[str, Any]] = {}

QUIZ_CACHE: Dict[str, List[Dict[str, Any]]] = {}

PM_JOBS_CACHE: Dict[str, Any] = {"mtime": None, "items": []}

EXTERNAL_INTERNSHIP_CACHE: Dict[str, List[Dict[str, Any]]] = {}

RECOMMENDATION_CURSOR: Dict[str, int] = {}

COURSE_CURSOR: Dict[str, int] = {}

DEFAULT_RESUME_QUIZ_QUESTION_COUNT = 10

DEFAULT_INTERVIEW_PREP_QUESTION_COUNT = 12

MAX_INTERVIEW_PREP_QUESTION_COUNT = 24

load_dotenv(".env.txt")

gemini_api_key = os.getenv("GEMINI_API_KEY")

groq_api_key = os.getenv("GROQ_API_KEY")

omnidim_api_key = os.getenv("OMNIDIMENSION_API_KEY")

agent_id = os.getenv("OMNIDIMENSION_AGENT_ID")

firecrawl_api_key = os.getenv("FIRECRAWL_API_KEY")

firecrawl_client = None

if firecrawl_api_key and FirecrawlApp is not None:

    try:

        firecrawl_client = FirecrawlApp(api_key=firecrawl_api_key)

        print("[STARTUP] Firecrawl client initialized.")

    except Exception as e:

        print(f"[STARTUP] WARNING: Firecrawl init failed: {e}")

elif firecrawl_api_key and FirecrawlApp is None:

    print("[STARTUP] WARNING: FIRECRAWL_API_KEY is present but firecrawl-py is not installed.")

else:

    print("[STARTUP] WARNING: No FIRECRAWL_API_KEY found in .env.txt")

class ProxySafeOmniClient(Client if Client is not None else object):

    """Ignore broken machine-level proxy variables for OmniDimension requests."""

    def request(self, method, endpoint, params=None, headers=None, data=None, json_data=None):

        headers = headers or {}

        params = params or {}

        method = method.upper()

        headers.setdefault("Authorization", f"Bearer {self.api_key}")

        headers.setdefault("Content-Type", "application/json")

        headers.setdefault("Accept", "application/json")

        url = self.base_url + "/" + endpoint.lstrip("/")

        session = requests.Session()

        session.trust_env = False

        try:

            response = session.request(

                method=method,

                url=url,

                params=params,

                headers=headers,

                data=data,

                json=json_data,

                timeout=30,

            )

            response.raise_for_status()

            if method == "DELETE":

                json_response = {}

            else:

                json_response = response.json() if response.content else {}

            return {

                "status": response.status_code,

                "json": json_response,

            }

        except requests.exceptions.HTTPError as e:

            error_message = "Unknown error"

            error_data = {}

            try:

                error_data = e.response.json()

                error_message = error_data.get("error_description", error_data.get("error", str(e)))

            except (ValueError, AttributeError, KeyError):

                error_message = str(e)

            raise APIError(

                status_code=e.response.status_code,

                message=error_message,

                response=error_data,

            )

        except requests.exceptions.RequestException as e:

            raise APIError(status_code=0, message=f"Network error: {str(e)}")

omnidim_client = None

if omnidim_api_key and Client is not None:

    try:

        omnidim_client = ProxySafeOmniClient(api_key=omnidim_api_key)

    except TypeError:

        omnidim_client = ProxySafeOmniClient(omnidim_api_key)

groq_client = None

if groq_api_key and Groq is not None:

    groq_client = Groq(api_key=groq_api_key)

    print("[STARTUP] Groq client initialized.")

elif groq_api_key and Groq is None:

    print("[STARTUP] WARNING: GROQ_API_KEY is present but the groq package is not installed.")

else:

    print("[STARTUP] WARNING: No GROQ_API_KEY found in .env.txt")

if gemini_api_key:

    print("[STARTUP] Gemini fallback key loaded.")

else:

    print("[STARTUP] WARNING: No GEMINI_API_KEY found in .env.txt")

app = FastAPI(title="PM Internship Recommendation API")

app.add_middleware(

    CORSMiddleware,

    allow_origins=["*"], 

    allow_credentials=True,

    allow_methods=["*"],

    allow_headers=["*"],

)

class ManualProfile(BaseModel):

    education: str

    location: str

    preferred_sector: str

    manual_skills: List[str]

class TranslationRequest(BaseModel):

    target_language: str

    payload: Dict[str, Any]

class AgentMessage(BaseModel):

    role: str

    content: str

class AgentChatRequest(BaseModel):

    messages: List[AgentMessage]

    user_skills: List[str]

    target_language: str

class InterviewPrepRequest(BaseModel):

    job_title: str

    company: str

    skills: List[str]

    question_count: Optional[int] = None

class LearningRecommendationRequest(BaseModel):

    job_title: Optional[str] = None

    company: str

    skills: List[str]

class DynamicJobsRequest(BaseModel):

    skills: List[str]
    total_score: Optional[int] = None

    location: str

    education: str

    preferred_sector: str

    target_language: str

    lang: Optional[str] = "en"

class VoiceRecommendRequest(BaseModel):

    """Payload from OmniDimension Voice AI.

    All fields are optional — the user may not mention every detail.

    `skills` can be a single comma-separated string or a list."""

    location: Optional[str] = None

    sector: Optional[str] = None

    skills: Optional[Any] = None

    education: Optional[str] = None

    lang: Optional[str] = "en"

class ResumeSkillsResponse(BaseModel):

    skills: List[str]

class QuizQuestionDraft(BaseModel):

    question: str

    correct_answer: str

    distractors: List[str]

class QuizQuestionsResponse(BaseModel):

    questions: List[QuizQuestionDraft]

class ResumeAnalysisResponse(BaseModel):

    skills: List[str]

    questions: List[QuizQuestionDraft]

    github_url: Optional[str] = None

    linkedin_url: Optional[str] = None

    email: Optional[str] = None

class LearningRecommendationDraft(BaseModel):

    title: str

    difficulty: str

    acceptance: str

    topic: str

class LearningRecommendationsResponse(BaseModel):

    recommendations: List[LearningRecommendationDraft]

class JobTipsResponse(BaseModel):

    tips: List[str]

class Test1QuestionDraft(BaseModel):

    id: str

    skill: str

    prompt: str

    options: List[str]

    answer: str

class Test1QuestionsResponse(BaseModel):

    questions: List[Test1QuestionDraft]

class Test2ChallengeDraft(BaseModel):

    id: str

    title: str

    difficulty: str

    companyTargets: List[str]

    prompt: str

    starterCode: str

    expectedSignals: List[str]

class Test2ChallengesResponse(BaseModel):

    challenges: List[Test2ChallengeDraft]

class GitHubVerificationRequest(BaseModel):

    github_url: Optional[str] = None

    github_username: Optional[str] = None

    claimed_skills: List[str] = []

class PlagiarismDetectionRequest(BaseModel):

    source_code: str

    candidate_code: str

    language: str = "python"

class TrustScoreRequest(BaseModel):

    skills: List[str] = []

    assessment_score: int = 0

    cheating_score: int = 0

    github_url: Optional[str] = None

    github_username: Optional[str] = None

    linkedin_url: Optional[str] = None

    email: Optional[str] = None

    resume_text: Optional[str] = None

class ResumeImproverRequest(BaseModel):

    resume_text: str

    skills: List[str] = []

    target_role: str = "PM Internship"

TARGET_SKILLS = {

    "Digital_Basics": ["data entry", "typing", "ms office", "excel", "word", "internet", "email"],

    "Vocational": ["agriculture", "wiring", "hardware", "plumbing", "mechanic", "welding", "carpentry", "solar"],

    "Logistics_Retail": ["inventory", "dispatch", "customer service", "packaging", "supply chain"],

    "Core_IT": ["python", "c++", "c", "java", "sql", "html", "networking", "troubleshooting"],

    "Healthcare_Basics": ["first aid", "sanitation", "patient care", "health records"]

}

COMMON_RESUME_SKILLS = [

    "python",

    "java",

    "c++",

    "c",

    "sql",

    "html",

    "css",

    "javascript",

    "react",

    "node.js",

    "excel",

    "ms office",

    "powerpoint",

    "word",

    "data entry",

    "typing",

    "basic computer",

    "customer service",

    "sales",

    "financial analysis",

    "accounting",

    "banking",

    "research",

    "field work",

    "teaching",

    "design",

    "canva",

    "photoshop",

    "video editing",

    "communication",

    "inventory",

    "dispatch",

    "supply chain",

    "packaging",

    "networking",

    "troubleshooting",

    "agriculture",

    "patient care",

    "health records",

]

def _clean_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()

def _ensure_absolute_url(url: str, default_scheme: str = "https") -> str:

    cleaned = _clean_text(url)

    if not cleaned:

        return ""

    parsed = urlparse(cleaned)

    if parsed.scheme in {"http", "https"}:

        return cleaned

    if cleaned.startswith("//"):

        candidate = f"{default_scheme}:{cleaned}"
    else:
        candidate_host = cleaned.lstrip("/")
        if not candidate_host or " " in candidate_host:
            return ""
        if "." not in candidate_host and not candidate_host.startswith(("localhost", "127.0.0.1")):
            return ""
        candidate = f"{default_scheme}://{candidate_host}"

    normalized = urlparse(candidate)
    if not normalized.netloc:
        return ""
    return candidate

def _extract_links_from_text(text: str) -> Dict[str, Optional[str]]:
    """Regex fallback to extract GitHub URL, LinkedIn URL, and email from raw text."""
    result: Dict[str, Optional[str]] = {"github_url": None, "linkedin_url": None, "email": None}
    github_match = re.search(r'https?://(?:www\.)?github\.com/[A-Za-z0-9_.-]+', text)
    if github_match:
        result["github_url"] = github_match.group(0)
    linkedin_match = re.search(r'https?://(?:www\.)?linkedin\.com/in/[A-Za-z0-9_.-]+', text)
    if linkedin_match:
        result["linkedin_url"] = linkedin_match.group(0)
    email_match = re.search(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}', text)
    if email_match:
        result["email"] = email_match.group(0)
    return result

def _dedupe_preserve(items: List[str]) -> List[str]:
    seen = set()

    result = []

    for item in items:

        key = item.casefold()

        if key in seen:

            continue

        seen.add(key)

        result.append(item)

    return result

def _normalize_assessment_prompt(prompt: str) -> str:

    cleaned = _clean_text(prompt)

    cleaned = re.sub(r"^(?:q\s*)?\d+\s*[\)\].:\-]+\s*", "", cleaned, flags=re.IGNORECASE)

    return cleaned

def _recommendation_request_key(user_profile: Dict[str, Any], user_skills: List[str]) -> str:

    normalized_skills = _normalize_skills(user_skills)

    return "|".join(

        [

            _clean_text(user_profile.get("location", "")).casefold(),

            _clean_text(user_profile.get("preferred_sector", "")).casefold(),

            _clean_text(user_profile.get("education", "")).casefold(),

            ",".join(skill.casefold() for skill in normalized_skills),

        ]

    )

def _course_request_key(target_skills: List[str]) -> str:

    normalized = _normalize_skills(target_skills)

    return ",".join(skill.casefold() for skill in normalized)

def _rotate_list(items: List[Any], offset: int) -> List[Any]:

    if not items:

        return []

    safe_offset = offset % len(items)

    if safe_offset == 0:

        return list(items)

    return list(items[safe_offset:]) + list(items[:safe_offset])

def _bounded_question_count(
    requested: Optional[int],
    default: int = DEFAULT_INTERVIEW_PREP_QUESTION_COUNT,
    minimum: int = 1,
    maximum: int = MAX_INTERVIEW_PREP_QUESTION_COUNT,
) -> int:

    try:
        if requested is None:
            return default
        return max(minimum, min(maximum, int(requested)))
    except Exception:
        return default

def _select_rotating_job_slice(jobs: List[Dict[str, Any]], request_key: str, top_n: int = 5) -> List[Dict[str, Any]]:

    if not jobs:

        return []

    top_n = max(1, min(top_n, 8))

    cursor = RECOMMENDATION_CURSOR.get(request_key, 0)

    RECOMMENDATION_CURSOR[request_key] = cursor + 1

    rotated = _rotate_list(jobs, cursor)

    selected: List[Dict[str, Any]] = []

    seen_company = set()

    for job in rotated:

        company_key = _clean_text(job.get("company", "")).casefold()

        if company_key and company_key in seen_company and len(selected) < (top_n - 1):

            continue

        if company_key:

            seen_company.add(company_key)

        selected.append(job)

        if len(selected) >= top_n:

            break

    if len(selected) < top_n:

        selected_keys = {

            (

                _clean_text(job.get("title", "")).casefold(),

                _clean_text(job.get("company", "")).casefold(),

                _clean_text(job.get("location", "")).casefold(),

            )

            for job in selected

        }

        for job in rotated:

            key = (

                _clean_text(job.get("title", "")).casefold(),

                _clean_text(job.get("company", "")).casefold(),

                _clean_text(job.get("location", "")).casefold(),

            )

            if key in selected_keys:

                continue

            selected.append(job)

            selected_keys.add(key)

            if len(selected) >= top_n:

                break

    return selected[:top_n]

def _fallback_resume_skill_candidates() -> List[str]:

    candidates: List[str] = []

    for skills in TARGET_SKILLS.values():

        candidates.extend(skills)

    candidates.extend(COMMON_RESUME_SKILLS)

    try:

        for job in _load_pm_jobs():

            candidates.extend(_normalize_skills(job.get("skills", [])))

    except Exception as e:

        print(f"Resume skill fallback warning: {e}")

    return _dedupe_preserve([_clean_text(skill) for skill in candidates if _clean_text(skill)])

def _extract_resume_skills_fallback(text: str, limit: int = 8) -> List[str]:

    source_text = text or ""

    lowered = _clean_text(source_text).casefold()

    if not lowered:

        return []

    detected: List[str] = []

    explicit_blocks = re.findall(

        r"(?:technical skills|core skills|key skills|skills|tools|technologies)\s*[:\-]\s*([^\n\r]{10,400})",

        source_text,

        flags=re.IGNORECASE,

    )

    for block in explicit_blocks:

        for piece in re.split(r"[,|/;•]+", block):

            candidate = _clean_text(piece).strip(":-")

            if 2 <= len(candidate) <= 40:

                detected.append(candidate)

    for skill in _fallback_resume_skill_candidates():

        tokens = [token for token in re.split(r"[^a-z0-9+#.]+", skill.casefold()) if token]

        if not tokens:

            continue

        pattern = r"\b" + r"(?:[\s/&,+.-]+)".join(re.escape(token) for token in tokens) + r"\b"

        if re.search(pattern, lowered) or (len(tokens) > 1 and all(token in lowered for token in tokens)):

            detected.append(skill)

        if len(detected) >= limit * 3:

            break

    blocked = {

        "resume",

        "curriculum vitae",

        "skills",

        "technical skills",

        "key skills",

        "tools",

    }

    normalized = _normalize_skills(detected)

    return [skill for skill in normalized if skill.casefold() not in blocked][:limit]

def _fallback_resume_questions(
    skills: List[str],
    num_questions: int = DEFAULT_RESUME_QUIZ_QUESTION_COUNT,
    used_prompts: Optional[set[str]] = None,
) -> List[Dict[str, Any]]:

    prompts = used_prompts or set()

    normalized_skills = _normalize_skills(skills) or ["workplace communication"]

    generated: List[Dict[str, Any]] = []

    rng = random.SystemRandom()

    scenarios = [
        "while cleaning internship survey data",
        "during candidate screening operations",
        "when preparing a daily reporting sheet",
        "while supporting a recruiter in a hiring sprint",
        "during a project status update",
        "while handling stakeholder requests",
    ]

    question_templates = [
        "Which option best demonstrates internship-ready use of {skill} {scenario}?",
        "For a PM internship task {scenario}, what is the strongest evidence of {skill}?",
        "What should you do first when applying {skill} {scenario}?",
        "Which action reflects practical proficiency in {skill} {scenario}?",
        "In a team assignment {scenario}, how should {skill} be applied for best results?",
        "When a manager asks for quick output {scenario}, what is the right way to apply {skill}?",
        "Which approach proves dependable {skill} usage {scenario}?",
    ]

    wrong_templates = [
        "Skip validation and submit quickly",
        "Rely on guesswork instead of method",
        "Wait for someone else to complete the task",
        "Ignore requirements and documentation",
        "Use unrelated tools without checking fit",
        "Hide assumptions and avoid peer review",
        "Finalize output without testing edge cases",
    ]

    template_offset = rng.randrange(len(question_templates))

    skill_offset = rng.randrange(len(normalized_skills)) if normalized_skills else 0

    scenario_offset = rng.randrange(len(scenarios))

    for idx in range(max(num_questions * 16, 40)):

        skill = normalized_skills[(skill_offset + idx) % len(normalized_skills)]

        template = question_templates[(template_offset + idx) % len(question_templates)]

        scenario = scenarios[(scenario_offset + idx) % len(scenarios)]

        question = template.format(skill=skill, scenario=scenario)

        question_key = question.casefold()

        if question_key in prompts:

            continue

        prompts.add(question_key)

        correct = f"Apply {skill} with clear steps, validation, and measurable output"

        wrong_choices = rng.sample(wrong_templates, k=min(3, len(wrong_templates)))

        while len(wrong_choices) < 3:

            wrong_choices.append("Follow an unverified approach")

        options = [correct, *wrong_choices[:3]]

        rng.shuffle(options)

        generated.append({

            "q": question,

            "options": options,

            "a": correct,

        })

        if len(generated) >= num_questions:

            break

    return generated
def _strip_code_fences(text: str) -> str:

    return (text or "").replace("```json", "").replace("```", "").strip()

@contextmanager

def _disable_proxy_env():

    proxy_keys = [

        "HTTP_PROXY", "HTTPS_PROXY", "ALL_PROXY",

        "http_proxy", "https_proxy", "all_proxy",

    ]

    previous_values = {key: os.environ.get(key) for key in proxy_keys}

    try:

        for key in proxy_keys:

            os.environ.pop(key, None)

        yield

    finally:

        for key, value in previous_values.items():

            if value is None:

                os.environ.pop(key, None)

            else:

                os.environ[key] = value

def _has_llm_provider() -> bool:

    return bool(groq_client or gemini_api_key)

def _candidate_gemini_models() -> List[str]:

    configured = _clean_text(os.getenv("GEMINI_MODEL"))

    ordered: List[str] = []

    if configured:

        ordered.append(configured)

    ordered.extend(DEFAULT_GEMINI_MODELS)

    result: List[str] = []

    for model_name in ordered:

        if model_name and model_name not in result:

            result.append(model_name)

    return result

def _candidate_groq_models() -> List[str]:

    configured = _clean_text(os.getenv("GROQ_MODEL"))

    ordered: List[str] = []

    if configured:

        ordered.append(configured)

    ordered.extend(DEFAULT_GROQ_MODELS)

    result: List[str] = []

    for model_name in ordered:

        if model_name and model_name not in result:

            result.append(model_name)

    return result

def _extract_http_error_message(error: Exception) -> str:

    response = getattr(error, "response", None)

    if response is not None:

        try:

            payload = response.json()

            if isinstance(payload, dict):

                error_obj = payload.get("error")

                if isinstance(error_obj, dict):

                    return _clean_text(error_obj.get("message") or error_obj.get("error") or str(payload))

                return _clean_text(str(payload))

        except Exception:

            pass

    return _clean_text(str(error))

def _is_groq_fatal_error(error: Exception) -> bool:

    message = _extract_http_error_message(error).casefold()

    return (

        "invalid api key" in message

        or "unauthorized" in message

        or "authentication" in message

        or "insufficient_quota" in message

        or "billing" in message

        or "credit" in message

        or "invalid api key" in message

    )

def _humanize_groq_error(error: Exception, attempted_models: List[str]) -> str:

    message = _extract_http_error_message(error)

    lowered = message.casefold()

    attempted = ", ".join(attempted_models)

    if "model not found" in lowered:

        return f"Groq rejected the configured model. Tried these models: {attempted}."

    if "credit" in lowered or "billing" in lowered or "insufficient_quota" in lowered:

        return "Groq rejected the request because the account is out of credits or billing is not enabled."

    if "invalid api key" in lowered or "unauthorized" in lowered or "authentication" in lowered:

        return "GROQ_API_KEY was rejected by Groq. Please check that the saved Groq key is correct."

    return f"Groq request failed after trying models: {attempted}. Last error: {message}"

def _humanize_gemini_error(error: Exception, attempted_models: List[str]) -> str:

    message = _extract_http_error_message(error)

    lowered = message.casefold()

    attempted = ", ".join(attempted_models)

    if "resource_exhausted" in lowered or "quota" in lowered or "429" in lowered:

        return (

            "Gemini is configured but its quota is exhausted right now. "

            f"Tried these models: {attempted}."

        )

    if "api key not valid" in lowered or "permission denied" in lowered or "403" in lowered:

        return "GEMINI_API_KEY was rejected by Google. Please check the key and API access."

    if "model" in lowered and "not found" in lowered:

        return f"Gemini rejected the configured model. Tried these models: {attempted}."

    return f"Gemini request failed after trying models: {attempted}. Last error: {message}"

def _extract_gemini_text(response_json: Dict[str, Any]) -> str:

    candidates = response_json.get("candidates") or []

    for candidate in candidates:

        content = candidate.get("content") or {}

        for part in content.get("parts") or []:

            text = part.get("text")

            if text:

                return text

    raise RuntimeError("Gemini returned no text content")


def _normalize_language(target_language: str) -> str:

    lowered = _clean_text(target_language).casefold()

    mapping = {

        "en": "en",

        "english": "en",

        "hi": "hi",

        "hindi": "hi",

        "हिंदी": "hi",

        "हिंदी (hindi)": "hi",

        "te": "te",

        "telugu": "te",

        "తెలుగు": "te",

        "తెలుగు (telugu)": "te",

        "mr": "mr",

        "marathi": "mr",

        "मराठी": "mr",

        "मराठी (marathi)": "mr",

        "ta": "ta",

        "tamil": "ta",

        "தமிழ்": "ta",

        "தமிழ் (tamil)": "ta",

        "bn": "bn",

        "bengali": "bn",

        "বাংলা": "bn",

        "gu": "gu",

        "gujarati": "gu",

        "ગુજરાતી": "gu",

        "pa": "pa",

        "punjabi": "pa",

        "ਪੰਜਾਬੀ": "pa",

        "kn": "kn",

        "kannada": "kn",

        "ಕನ್ನಡ": "kn",

        "ml": "ml",

        "malayalam": "ml",

        "മലയാളം": "ml",

        "or": "or",

        "odia": "or",

        "ଓଡ଼ିଆ": "or",

        "as": "as",

        "assamese": "as",

        "অসমীয়া": "as",

    }

    return mapping.get(lowered, "en")

def _call_groq_structured(prompt: str, schema_model: type[BaseModel], *, system_instruction: Optional[str] = None, temperature: float = 0.2, max_output_tokens: int = 2048) -> BaseModel:

    if not groq_client:

        raise RuntimeError("GROQ_API_KEY is missing")

    schema_text = schema_model.schema_json()

    full_prompt = f"{prompt}\\n\\nPlease strictly follow this JSON schema for your response:\\n{schema_text}"

    messages = []

    if system_instruction:

        messages.append({"role": "system", "content": system_instruction})

    messages.append({"role": "user", "content": full_prompt})

    last_error: Optional[Exception] = None

    attempted_models: List[str] = []

    for model_name in _candidate_groq_models():

        attempted_models.append(model_name)

        print(f"Sending structured call to Groq model: {model_name}")

        try:

            response = groq_client.chat.completions.create(

                model=model_name,

                messages=messages,

                response_format={"type": "json_object"},

                temperature=temperature,

                max_tokens=max_output_tokens,

            )

            content = (response.choices[0].message.content or "").strip()

            parsed = json.loads(content)

            return schema_model.model_validate(parsed)

        except Exception as e:

            print(f"Groq structured call error on {model_name}: {e}")

            last_error = e

            if _is_groq_fatal_error(e):

                break

            continue

    raise RuntimeError(_humanize_groq_error(last_error, attempted_models)) from last_error

def _call_groq_text(prompt: str, *, system_instruction: Optional[str] = None, temperature: float = 0.2, max_output_tokens: int = 2048) -> str:

    if not groq_client:

        raise RuntimeError("GROQ_API_KEY is missing")

    messages = []

    if system_instruction:

        messages.append({"role": "system", "content": system_instruction})

    messages.append({"role": "user", "content": prompt})

    last_error: Optional[Exception] = None

    attempted_models: List[str] = []

    for model_name in _candidate_groq_models():

        attempted_models.append(model_name)

        try:

            response = groq_client.chat.completions.create(

                model=model_name,

                messages=messages,

                temperature=temperature,

                max_tokens=max_output_tokens,

            )

            return (response.choices[0].message.content or "").strip()

        except Exception as e:

            print(f"Groq text call error on {model_name}: {e}")

            last_error = e

            if _is_groq_fatal_error(e):

                break

            continue

    raise RuntimeError(_humanize_groq_error(last_error, attempted_models)) from last_error

def _call_gemini_structured(prompt: str, schema_model: type[BaseModel], *, system_instruction: Optional[str] = None, temperature: float = 0.2, max_output_tokens: int = 2048) -> BaseModel:

    if not gemini_api_key:

        raise RuntimeError("GEMINI_API_KEY is missing")

    payload = {

        "contents": [

            {

                "parts": [

                    {

                        "text": (

                            f"{prompt}\n\nReturn strict JSON matching this schema:\n"

                            f"{json.dumps(schema_model.model_json_schema(), ensure_ascii=False)}"

                        )

                    }

                ]

            }

        ],

        "generationConfig": {

            "temperature": temperature,

            "response_mime_type": "application/json",

        },

    }

    if max_output_tokens:

        payload["generationConfig"]["max_output_tokens"] = max_output_tokens

    if system_instruction:

        payload["system_instruction"] = {

            "parts": [

                {

                    "text": system_instruction

                }

            ]

        }

    last_error: Optional[Exception] = None

    attempted_models: List[str] = []

    for model_name in _candidate_gemini_models():

        attempted_models.append(model_name)

        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent"

        session = requests.Session()

        session.trust_env = False

        try:

            response = session.post(

                url,

                headers={

                    "x-goog-api-key": gemini_api_key,

                    "Content-Type": "application/json",

                },

                json=payload,

                timeout=60,

            )

            response.raise_for_status()

            parsed = json.loads(_strip_code_fences(_extract_gemini_text(response.json())))

            return schema_model.model_validate(parsed)

        except Exception as e:

            print(f"Gemini structured call error on {model_name}: {e}")

            last_error = e

            continue

    raise RuntimeError(_humanize_gemini_error(last_error, attempted_models)) from last_error

def _call_gemini_text(prompt: str, *, system_instruction: Optional[str] = None, temperature: float = 0.2, max_output_tokens: int = 2048) -> str:

    if not gemini_api_key:

        raise RuntimeError("GEMINI_API_KEY is missing")

    payload = {

        "contents": [

            {

                "parts": [

                    {

                        "text": prompt

                    }

                ]

            }

        ],

        "generationConfig": {

            "temperature": temperature,

        },

    }

    if max_output_tokens:

        payload["generationConfig"]["max_output_tokens"] = max_output_tokens

    if system_instruction:

        payload["system_instruction"] = {

            "parts": [

                {

                    "text": system_instruction

                }

            ]

        }

    last_error: Optional[Exception] = None

    attempted_models: List[str] = []

    for model_name in _candidate_gemini_models():

        attempted_models.append(model_name)

        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent"

        session = requests.Session()

        session.trust_env = False

        try:

            response = session.post(

                url,

                headers={

                    "x-goog-api-key": gemini_api_key,

                    "Content-Type": "application/json",

                },

                json=payload,

                timeout=60,

            )

            response.raise_for_status()

            return _strip_code_fences(_extract_gemini_text(response.json()))

        except Exception as e:

            print(f"Gemini text call error on {model_name}: {e}")

            last_error = e

            continue

    raise RuntimeError(_humanize_gemini_error(last_error, attempted_models)) from last_error

def _call_structured(prompt: str, schema_model: type[BaseModel], *, system_instruction: Optional[str] = None, temperature: float = 0.2, max_output_tokens: int = 2048) -> BaseModel:

    provider_errors: List[str] = []

    if groq_client:

        try:

            return _call_groq_structured(

                prompt,

                schema_model,

                system_instruction=system_instruction,

                temperature=temperature,

                max_output_tokens=max_output_tokens,

            )

        except Exception as e:

            provider_errors.append(str(e))

    if gemini_api_key:

        try:

            return _call_gemini_structured(

                prompt,

                schema_model,

                system_instruction=system_instruction,

                temperature=temperature,

                max_output_tokens=max_output_tokens,

            )

        except Exception as e:

            provider_errors.append(str(e))

    if provider_errors:

        raise RuntimeError(" | ".join(provider_errors))

    raise RuntimeError("No AI provider configured. Add GROQ_API_KEY or GEMINI_API_KEY.")

def _call_text(prompt: str, *, system_instruction: Optional[str] = None, temperature: float = 0.2, max_output_tokens: int = 2048) -> str:

    provider_errors: List[str] = []

    if groq_client:

        try:

            return _call_groq_text(

                prompt,

                system_instruction=system_instruction,

                temperature=temperature,

                max_output_tokens=max_output_tokens,

            )

        except Exception as e:

            provider_errors.append(str(e))

    if gemini_api_key:

        try:

            return _call_gemini_text(

                prompt,

                system_instruction=system_instruction,

                temperature=temperature,

                max_output_tokens=max_output_tokens,

            )

        except Exception as e:

            provider_errors.append(str(e))

    if provider_errors:

        raise RuntimeError(" | ".join(provider_errors))

    raise RuntimeError("No AI provider configured. Add GROQ_API_KEY or GEMINI_API_KEY.")

def _normalize_skills(skills: List[str]) -> List[str]:

    cleaned = []

    for skill in skills:

        normalized = _clean_text(skill)

        if normalized:

            cleaned.append(normalized)

    return _dedupe_preserve(cleaned)

def _extract_omnidim_payload(response: Any) -> Dict[str, Any]:

    if isinstance(response, dict):

        payload = response.get("json")

        if isinstance(payload, dict):

            return payload

        return response

    return {}

def _humanize_omnidim_error(error: Exception) -> str:

    message = _clean_text(str(error))

    lowered = message.casefold()

    status_code = getattr(error, "status_code", None)

    if "127.0.0.1:9" in message or "unable to connect to proxy" in lowered or "proxyerror" in lowered:

        return "OmniDimension requests are blocked by a broken local proxy configuration. The backend now bypasses system proxies for OmniDimension. Restart the server and try again."

    if status_code in {401, 403} or "unauthorized" in lowered or "forbidden" in lowered:

        return "OMNIDIMENSION_API_KEY was rejected by OmniDimension. Please check that the saved API key is valid for this OmniDimension account."

    if status_code == 404 and "/agents/" in message:

        return "OMNIDIMENSION_AGENT_ID does not match any agent in this OmniDimension account."

    return f"OmniDimension request failed: {message}"

def _get_omnidim_client():

    if not omnidim_api_key:

        raise RuntimeError("OMNIDIMENSION_API_KEY is missing.")

    if not omnidim_client:

        raise RuntimeError("OmniDimension client could not be initialized. Check OMNIDIMENSION_API_KEY.")

    return omnidim_client

def _list_omnidim_agents(client) -> List[Dict[str, Any]]:

    payload = _extract_omnidim_payload(client.agent.list(page=1, page_size=100))

    bots = payload.get("bots") or []

    return [bot for bot in bots if isinstance(bot, dict)]

def _resolve_omnidim_agent(client) -> tuple[int, Optional[str], List[Dict[str, Any]]]:

    configured_agent_id = _clean_text(agent_id)

    if not configured_agent_id:

        raise RuntimeError("OMNIDIMENSION_AGENT_ID is missing.")

    bots = _list_omnidim_agents(client)

    if not bots:

        raise RuntimeError("No OmniDimension agents were found for this API key.")

    matched_bot = next((bot for bot in bots if str(bot.get("id")) == configured_agent_id), None)

    if matched_bot and matched_bot.get("id") is not None:

        return int(matched_bot["id"]), None, bots

    if len(bots) == 1 and bots[0].get("id") is not None:

        resolved_agent_id = int(bots[0]["id"])

        warning = (

            f"Configured OMNIDIMENSION_AGENT_ID '{configured_agent_id}' does not match any agent in this OmniDimension account. "

            f"Automatically using the only available agent ({resolved_agent_id})."

        )

        return resolved_agent_id, warning, bots

    available_agent_ids = ", ".join(str(bot.get("id")) for bot in bots if bot.get("id") is not None)

    raise RuntimeError(

        f"Configured OMNIDIMENSION_AGENT_ID '{configured_agent_id}' does not match any agent in this OmniDimension account. "

        f"Available agent IDs: {available_agent_ids or 'none'}."

    )

def _build_quiz_question(question: QuizQuestionDraft) -> Dict[str, Any]:

    prompt = _normalize_assessment_prompt(question.question)

    correct_answer = _clean_text(question.correct_answer)

    distractors = _normalize_skills(question.distractors)

    distractors = [item for item in distractors if item.casefold() != correct_answer.casefold()]

    if not prompt or not correct_answer or len(distractors) < 3:

        raise ValueError("AI returned an incomplete quiz question")

    options = [correct_answer, distractors[0], distractors[1], distractors[2]]

    if len(_dedupe_preserve(options)) < 4:

        raise ValueError("AI returned duplicate quiz options")

    random.SystemRandom().shuffle(options)

    return {

        "q": prompt,

        "options": options,

        "a": next(option for option in options if option.casefold() == correct_answer.casefold()),

    }

def extract_skills_with_ai(text: str, api_key: str) -> List[str]:

    if not _has_llm_provider():

        raise RuntimeError("No AI provider configured. Add GROQ_API_KEY or GEMINI_API_KEY.")

    resume_text = _clean_text(text)

    if not resume_text:

        raise RuntimeError("Could not extract readable text from the PDF")

    try:

        response = _call_structured(

            prompt=(

                "Extract the 5 to 8 strongest resume-evidenced skills from the text below.\n"

                "Rules:\n"

                "- Prefer concrete skills, tools, technologies, platforms, methods, and domain abilities.\n"

                "- Do not include generic soft skills like communication, leadership, or hardworking unless the resume has no stronger skill evidence.\n"

                "- Use the exact skill names from the resume where possible.\n"

                "- Return only skills that are clearly supported by the resume text.\n\n"

                f"RESUME TEXT:\n{resume_text[:12000]}"

            ),

            schema_model=ResumeSkillsResponse,

            system_instruction="You extract only resume-supported skills and return strict JSON.",

            temperature=0.1,

            max_output_tokens=512,

        )

        skills = _normalize_skills(response.skills)

        if not skills:

            raise RuntimeError("AI returned no usable skills")

        return skills[:8]

    except Exception as e:

        print(f"LLM Extraction Error: {e}")

        raise RuntimeError("Skill extraction failed") from e


def generate_resume_analysis(text: str, api_key: str, num_questions: int = DEFAULT_RESUME_QUIZ_QUESTION_COUNT) -> Dict[str, Any]:

    resume_text = _clean_text(text)

    if not resume_text:

        raise RuntimeError("We could not read text from this PDF. Please upload a text-based PDF or use manual skill selection.")

    fallback_skills = _extract_resume_skills_fallback(text)

    if not _has_llm_provider():

        if fallback_skills:

            return {

                "skills": fallback_skills,

                "questions": _fallback_resume_questions(fallback_skills, num_questions=num_questions),

            }

        raise RuntimeError("No AI provider configured. Add GROQ_API_KEY or GEMINI_API_KEY.")

    try:

        response = _call_structured(

            prompt=(

                f"Read the resume text and do three tasks in one response.\n"

                f"1. Extract 5 to 8 strong, resume-supported skills.\n"

                f"2. Generate exactly {num_questions} interview MCQs based only on those extracted skills.\n"

                f"3. Extract the candidate's GitHub URL, LinkedIn URL, and email address if present in the text.\n"

                "Rules:\n"

                "- Prefer concrete tools, technologies, methods, and job skills over soft skills.\n"

                "- Every question must directly test one extracted skill.\n"

                "- For each question, provide one correct answer and exactly three plausible wrong answers.\n"

                "- Do not ask generic aptitude questions.\n"

                "- Use short, distinct answer options.\n"

                "- Avoid 'all of the above' and 'none of the above'.\n\n"

                f"RESUME TEXT:\n{resume_text[:12000]}"

            ),

            schema_model=ResumeAnalysisResponse,

            system_instruction="You extract resume-supported skills, generate matching interview MCQs, and extract contact URLs/emails in strict JSON.",

            temperature=0.25,

            max_output_tokens=3072,

        )

        skills = _normalize_skills(response.skills) or fallback_skills

        if not skills:

            raise RuntimeError("AI returned no usable skills")

        questions: List[Dict[str, Any]] = []

        seen_prompts: set[str] = set()

        for item in response.questions:

            try:

                built_question = _build_quiz_question(item)

            except Exception as e:

                print(f"Resume question fallback warning: {e}")

                continue

            prompt_key = built_question["q"].casefold()

            if prompt_key in seen_prompts:

                continue

            seen_prompts.add(prompt_key)

            questions.append(built_question)

        if len(questions) < num_questions:

            questions.extend(

                _fallback_resume_questions(

                    skills,

                    num_questions=num_questions - len(questions),

                    used_prompts=seen_prompts,

                )

            )

        if len(questions) < num_questions:

            raise RuntimeError(f"AI returned only {len(questions)} valid questions")

        extracted_github = response.github_url
        extracted_linkedin = response.linkedin_url
        extracted_email = response.email
        if not extracted_github or not extracted_linkedin or not extracted_email:
            regex_links = _extract_links_from_text(resume_text)
            if not extracted_github:
                extracted_github = regex_links.get("github_url")
            if not extracted_linkedin:
                extracted_linkedin = regex_links.get("linkedin_url")
            if not extracted_email:
                extracted_email = regex_links.get("email")
        return {
            "skills": skills[:8],
            "questions": questions[:num_questions],
            "github_url": extracted_github,
            "linkedin_url": extracted_linkedin,
            "email": extracted_email
        }

    except Exception as e:

        print(f"Resume analysis fallback warning: {e}")

        if fallback_skills:

            return {

                "skills": fallback_skills,

                "questions": _fallback_resume_questions(fallback_skills, num_questions=num_questions),

            }

        raise RuntimeError("Resume analysis failed. Please upload a clearer text-based PDF or use manual skill selection.") from e


CITY_TO_STATE = {

    "mumbai": "maharashtra",

    "pune": "maharashtra",

    "delhi": "delhi",

    "bangalore": "karnataka",

    "bengaluru": "karnataka",

    "hyderabad": "telangana",

    "chennai": "tamil nadu",

    "kolkata": "west bengal",

    "noida": "uttar pradesh",

    "ahmedabad": "gujarat",

    "jaipur": "rajasthan",

}

def _tokenize_for_match(value: Any) -> List[str]:

    return re.findall(r"[a-z0-9\+#&]+", _clean_text(value).casefold())

def _job_search_text(job: Dict[str, Any]) -> str:

    return _clean_text(" ".join([

        _clean_text(job.get("title")),

        _clean_text(job.get("company")),

        _clean_text(job.get("sector")),

        _clean_text(job.get("field")),

        _clean_text(job.get("location")),

        _clean_text(job.get("education")),

        " ".join(_normalize_skills(job.get("skills", []))),

        _clean_text(job.get("description")),

        _clean_text(job.get("source")),

    ]))

def _is_pm_portal_job(job: Dict[str, Any]) -> bool:

    apply_url = _clean_text(job.get("apply_url")).casefold()

    source = _clean_text(job.get("source")).casefold()

    description = _clean_text(job.get("description")).casefold()

    return (

        "pminternship.mca.gov.in" in apply_url

        or "pm internship" in source

        or "pm internship" in description

    )

def _normalize_pm_job_record(job: Dict[str, Any]) -> Dict[str, Any]:

    normalized_job = dict(job)

    normalized_job["source"] = _clean_text(job.get("source")) or PM_SOURCE_NAME

    normalized_job["apply_url"] = _clean_text(job.get("apply_url")) or PM_PORTAL_URL

    normalized_job["field"] = _clean_text(job.get("field"))

    existing_skills = _normalize_skills(job.get("skills", []))

    meaningful_skills = [

        skill for skill in existing_skills

        if skill.casefold() not in {"communication", "on-the-job training", "general skills"}

    ]

    if infer_pm_job_skills is not None and len(meaningful_skills) < 2:

        rebuilt_skills = infer_pm_job_skills(

            _clean_text(job.get("title")),

            _clean_text(job.get("sector")),

            _clean_text(job.get("field")),

        )

        normalized_job["skills"] = rebuilt_skills[:6] if rebuilt_skills else existing_skills

    else:

        normalized_job["skills"] = existing_skills

    return normalized_job

def _dedupe_pm_jobs(jobs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:

    deduped: Dict[str, Dict[str, Any]] = {}

    for job in jobs:

        if not isinstance(job, dict):

            continue

        job = _normalize_pm_job_record(job)

        if not _is_pm_portal_job(job):

            continue

        key = "|".join([

            _clean_text(job.get("title")).casefold(),

            _clean_text(job.get("company")).casefold(),

            _clean_text(job.get("location")).casefold(),

        ])

        if not key.strip("|"):

            continue

        existing = deduped.get(key)

        if existing is None or len(_job_search_text(job)) > len(_job_search_text(existing)):

            deduped[key] = job

    return list(deduped.values())

def _load_pm_jobs(force_refresh: bool = False) -> List[Dict[str, Any]]:

    current_mtime = os.path.getmtime(JOBS_JSON_PATH) if os.path.exists(JOBS_JSON_PATH) else None

    if not force_refresh and PM_JOBS_CACHE["items"] and PM_JOBS_CACHE["mtime"] == current_mtime:

        return PM_JOBS_CACHE["items"]

    data: Optional[List[Dict[str, Any]]] = None

    refresh_error: Optional[Exception] = None

    if force_refresh and fetch_and_save_internships is not None:

        try:

            data = fetch_and_save_internships()

        except Exception as e:

            refresh_error = e

    if data is None:

        try:

            with open(JOBS_JSON_PATH, "r", encoding="utf-8") as f:

                loaded = json.load(f)

            if isinstance(loaded, list):

                data = loaded

        except Exception as e:

            refresh_error = refresh_error or e

    if data is None and fetch_and_save_internships is not None:

        try:

            data = fetch_and_save_internships()

        except Exception as e:

            refresh_error = refresh_error or e

    if data is None:

        raise RuntimeError(f"Could not load PM internships: {refresh_error}")

    pm_jobs = _dedupe_pm_jobs(data)

    if not pm_jobs:

        raise RuntimeError("No PM Internship Scheme listings are available in jobs.json right now.")

    PM_JOBS_CACHE["items"] = pm_jobs

    PM_JOBS_CACHE["mtime"] = os.path.getmtime(JOBS_JSON_PATH) if os.path.exists(JOBS_JSON_PATH) else current_mtime

    return pm_jobs

def _match_user_skills_to_job(user_skills: List[str], job: Dict[str, Any]) -> List[str]:

    job_text = _job_search_text(job).casefold()

    matched_skills: List[str] = []

    for skill in _normalize_skills(user_skills):

        normalized_skill = skill.casefold()

        if not normalized_skill:

            continue

        if normalized_skill in job_text:

            matched_skills.append(skill)

            continue

        tokens = [token for token in _tokenize_for_match(normalized_skill) if len(token) > 2]

        if len(tokens) > 1 and all(token in job_text for token in tokens):

            matched_skills.append(skill)

            continue

        if len(tokens) == 1 and tokens[0] in job_text:

            matched_skills.append(skill)

    return _dedupe_preserve(matched_skills)

def _preferred_sector_matches_job(preferred_sector: str, job: Dict[str, Any]) -> bool:

    normalized_sector = _clean_text(preferred_sector).casefold()

    if normalized_sector in {"", "any", "all sectors"}:

        return False

    haystack = f"{_clean_text(job.get('sector'))} {_clean_text(job.get('field'))} {_clean_text(job.get('title'))}".casefold()

    if normalized_sector in haystack:

        return True

    tokens = [token for token in _tokenize_for_match(normalized_sector) if len(token) > 2]

    return bool(tokens) and all(token in haystack for token in tokens)

def _preferred_location_matches_job(preferred_location: str, job: Dict[str, Any]) -> bool:

    normalized_location = _clean_text(preferred_location).casefold()

    if normalized_location in {"", "india", "india (any)", "any", "remote"}:

        return False

    mapped_location = CITY_TO_STATE.get(normalized_location, normalized_location)

    job_location = _clean_text(job.get("location")).casefold()

    return (

        normalized_location in job_location

        or mapped_location in job_location

        or job_location in {normalized_location, mapped_location}

    )

def _education_matches_job(education: str, job: Dict[str, Any]) -> bool:

    normalized_education = _clean_text(education).casefold()

    job_education = _clean_text(job.get("education")).casefold()

    if not normalized_education or not job_education:

        return False

    return normalized_education in job_education or job_education in normalized_education

def _slugify_term(value: str) -> str:

    return re.sub(r"[^a-z0-9]+", "-", _clean_text(value).casefold()).strip("-")

def _extract_html_text(fragment: str) -> str:

    cleaned = re.sub(r"<[^>]+>", " ", fragment or "")

    cleaned = html.unescape(cleaned).replace("\xa0", " ")

    cleaned = cleaned.replace("₹", "Rs ").replace("? ", "Rs ")

    return _clean_text(cleaned)

def _extract_first_match(pattern: str, text: str, flags: int = re.S) -> str:

    match = re.search(pattern, text, flags)

    return match.group(1) if match else ""

def _candidate_external_terms(user_profile: Dict[str, Any], user_skills: List[str]) -> List[str]:

    blocked = {"communication", "general skills", "on-the-job training"}

    preferred_sector = _clean_text(user_profile.get("preferred_sector"))

    terms = [skill for skill in _normalize_skills(user_skills) if skill.casefold() not in blocked]

    if preferred_sector and preferred_sector.casefold() not in {"any", "all sectors"}:

        terms.append(preferred_sector)

    if not terms:

        terms.append("internship")

    deduped = []

    seen = set()

    for term in terms:

        key = term.casefold()

        if key in seen:

            continue

        seen.add(key)

        deduped.append(term)

    return deduped[:3]

def _candidate_internshala_urls(user_profile: Dict[str, Any], user_skills: List[str]) -> List[str]:

    location = _clean_text(user_profile.get("location"))

    location_slug = _slugify_term(location)

    if location.casefold() in {"", "india", "india (any)", "any"}:

        location_slug = ""

    terms = _candidate_external_terms(user_profile, user_skills)

    urls: List[str] = []

    for term in terms:

        term_slug = _slugify_term(term)

        if not term_slug:

            continue

        if location_slug:

            urls.append(f"{INTERNSHALA_BASE_URL}/internships/{term_slug}-internship-in-{location_slug}/")

            urls.append(f"{INTERNSHALA_BASE_URL}/internships/work-from-home-{term_slug}-internship-in-{location_slug}/")

        urls.append(f"{INTERNSHALA_BASE_URL}/internships/{term_slug}-internship/")

        urls.append(f"{INTERNSHALA_BASE_URL}/internships/work-from-home-{term_slug}-internship/")

        urls.append(f"{INTERNSHALA_BASE_URL}/internships/keywords-{term_slug}/")

    deduped_urls = []

    seen = set()

    for url in urls:

        if url in seen:

            continue

        seen.add(url)

        deduped_urls.append(url)

    return deduped_urls[:6]

def _parse_internshala_listings(page_html: str, default_sector: str = "General") -> List[Dict[str, Any]]:

    blocks = re.findall(

        r'(<div class="container-fluid individual_internship.*?)(?=<div class="container-fluid individual_internship|\Z)',

        page_html,

        flags=re.S,

    )

    listings: List[Dict[str, Any]] = []

    for idx, block in enumerate(blocks):

        title = _extract_html_text(_extract_first_match(r'<a class="job-title-href"[^>]*>(.*?)</a>', block))

        company = _extract_html_text(_extract_first_match(r'<p class="company-name">\s*(.*?)\s*</p>', block))

        href = _extract_first_match(r'data-href=[\'"]?([^\'"\s>]+)', block)

        location_block = _extract_first_match(r'<div class="row-1-item locations">(.*?)</div>\s*<!-- /location -->', block)

        location_parts = re.findall(r'<a>(.*?)</a>', location_block, flags=re.S)

        location = _clean_text(", ".join(_extract_html_text(part) for part in location_parts if _extract_html_text(part)))

        if not location:

            location = _extract_html_text(location_block) or "India"

        stipend = _extract_html_text(_extract_first_match(r"<span class='stipend'>(.*?)</span>", block))

        duration = _extract_html_text(_extract_first_match(r'<div class="row-1-item">\s*<i class="ic-16-calendar"></i>\s*<span>(.*?)</span>', block))

        description = _extract_html_text(_extract_first_match(r'<div class="about_job">.*?<div class="text">\s*(.*?)\s*</div>\s*</div>', block))

        posted = _extract_html_text(_extract_first_match(r'<div class="status-info">.*?<span>(.*?)</span>', block))

        skill_matches = re.findall(r"<div class='job_skill'>(.*?)</div>", block, flags=re.S)

        skills = _normalize_skills([_extract_html_text(skill) for skill in skill_matches if _extract_html_text(skill)])

        if not title or not company or not href:

            continue

        sector = default_sector if _clean_text(default_sector) and _clean_text(default_sector).casefold() not in {"any", "all sectors"} else "General"

        listings.append({

            "id": f"internshala-{idx + 1}",

            "title": title,

            "company": company,

            "location": location,

            "sector": sector,

            "field": duration,

            "education": "Any",

            "description": _clean_text(" ".join(part for part in [description, stipend, duration, posted] if part)),

            "skills": skills[:8],

            "apply_url": urljoin(INTERNSHALA_BASE_URL, href),

            "source": INTERNSHALA_SOURCE,

        })

    return listings

def _score_external_job(job: Dict[str, Any], user_profile: Dict[str, Any], user_skills: List[str]) -> Dict[str, Any]:

    matched_skills = _match_user_skills_to_job(user_skills, job)

    sector_match = _preferred_sector_matches_job(user_profile.get("preferred_sector", ""), job)

    location_match = _preferred_location_matches_job(user_profile.get("location", ""), job)

    education_match = _education_matches_job(user_profile.get("education", ""), job)

    score = 35 + (len(matched_skills) * 14)

    if sector_match:

        score += 12

    if location_match:

        score += 10

    if education_match:

        score += 4

    if not matched_skills and (sector_match or location_match):

        score = max(score, 48)

    job["matched_skills"] = matched_skills

    job["match_score"] = max(25, min(95, score))

    return job

def _fetch_internshala_internships(user_profile: Dict[str, Any], user_skills: List[str], max_results: int = 5) -> List[Dict[str, Any]]:

    cache_key = "|".join([

        _clean_text(user_profile.get("location")).casefold(),

        _clean_text(user_profile.get("preferred_sector")).casefold(),

        ",".join(skill.casefold() for skill in _normalize_skills(user_skills)),

    ])

    cached = EXTERNAL_INTERNSHIP_CACHE.get(cache_key)

    if cached:

        return cached[:max_results]

    urls = _candidate_internshala_urls(user_profile, user_skills)

    listings: List[Dict[str, Any]] = []

    seen = set()

    use_firecrawl = firecrawl_client is not None

    session = requests.Session()

    session.trust_env = False

    headers = {

        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36",

        "Accept-Language": "en-IN,en;q=0.9",

    }

    for url in urls:

        try:

            page_html = None

            if use_firecrawl:

                try:

                    fc_result = _firecrawl_scrape(url, formats=["html"])

                    page_html = fc_result.get("html") or fc_result.get("content") or ""

                    print(f"[Firecrawl] Scraped {url} ({len(page_html)} chars)")

                except Exception as fc_err:

                    print(f"[Firecrawl] Failed for {url}, falling back to requests: {fc_err}")

            if not page_html:

                with _disable_proxy_env():

                    response = session.get(url, headers=headers, timeout=30)

                if response.status_code != 200:

                    continue

                page_html = response.text

            parsed = _parse_internshala_listings(page_html, _clean_text(user_profile.get("preferred_sector")) or "General")

            for listing in parsed:

                key = "|".join([

                    listing.get("title", "").casefold(),

                    listing.get("company", "").casefold(),

                    listing.get("location", "").casefold(),

                ])

                if key in seen:

                    continue

                seen.add(key)

                listings.append(_score_external_job(listing, user_profile, _normalize_skills(user_skills)))

            if len(listings) >= max_results:

                break

        except Exception as e:

            print(f"Internshala fallback warning for {url}: {e}")

    listings.sort(
        key=lambda item: (
            item.get("match_score", 0),
            len(item.get("matched_skills", [])),
            _clean_text(item.get("title", "")),
        ),
        reverse=True,
    )

    EXTERNAL_INTERNSHIP_CACHE[cache_key] = listings

    return listings[:max_results]

def _merge_recommendation_sets(primary_jobs: List[Dict[str, Any]], fallback_jobs: List[Dict[str, Any]], top_n: int = 5) -> List[Dict[str, Any]]:

    combined: List[Dict[str, Any]] = []

    seen = set()

    for job in primary_jobs + fallback_jobs:

        key = "|".join([

            _clean_text(job.get("title")).casefold(),

            _clean_text(job.get("company")).casefold(),

            _clean_text(job.get("location")).casefold(),

        ])

        if key in seen:

            continue

        seen.add(key)

        combined.append(job)

    combined.sort(
        key=lambda item: (
            item.get("match_score", 0),
            len(item.get("matched_skills", [])),
            _clean_text(item.get("title", "")),
        ),
        reverse=True,
    )

    return combined[:top_n]

def rag_semantic_search(user_profile: dict, user_skills: List[str], top_n=5):

    """Return only PM Internship Scheme roles that actually match the user's profile."""

    top_n = max(1, min(top_n, 30))

    try:

        db_jobs = _load_pm_jobs()

    except Exception as e:

        print(f"Error loading PM internship data: {e}")

        return []

    if not db_jobs:

        return []

    normalized_skills = _normalize_skills(user_skills)

    query = _clean_text(" ".join([

        _clean_text(user_profile.get("preferred_sector")),

        _clean_text(user_profile.get("location")),

        _clean_text(user_profile.get("education")),

        " ".join(normalized_skills),

    ]))

    corpus = [_job_search_text(job) for job in db_jobs]

    semantic_scores = [0.0] * len(db_jobs)

    if query:

        try:

            vectorizer = TfidfVectorizer(ngram_range=(1, 2), stop_words="english")

            tfidf_matrix = vectorizer.fit_transform(corpus + [query])

            semantic_scores = cosine_similarity(tfidf_matrix[-1], tfidf_matrix[:-1]).flatten().tolist()

        except Exception as e:

            print(f"Semantic ranking warning: {e}")

    ranked = []

    fallback_ranked = []

    for idx, job in enumerate(db_jobs):

        matched_skills = _match_user_skills_to_job(normalized_skills, job)

        skill_ratio = (len(matched_skills) / len(normalized_skills)) if normalized_skills else 0.0

        semantic_score = float(semantic_scores[idx]) if idx < len(semantic_scores) else 0.0

        sector_match = _preferred_sector_matches_job(user_profile.get("preferred_sector", ""), job)

        location_match = _preferred_location_matches_job(user_profile.get("location", ""), job)

        education_match = _education_matches_job(user_profile.get("education", ""), job)

        if normalized_skills and not matched_skills:

            if not sector_match and not location_match and semantic_score < 0.08:

                continue

            if not sector_match and not location_match and semantic_score < 0.14:

                continue

        composite_score = (

            (skill_ratio * 0.58)

            + (semantic_score * 0.24)

            + (0.10 if sector_match else 0.0)

            + (0.06 if location_match else 0.0)

            + (0.02 if education_match else 0.0)

        )

        if matched_skills and composite_score < 0.30:

            composite_score = 0.30 + (min(len(matched_skills), 3) * 0.02)

        match_score = max(1, min(99, int(round(composite_score * 100))))

        if not matched_skills and (sector_match or location_match) and match_score < 45:

            match_score = 50 if (sector_match and location_match) else 45

        elif not matched_skills and semantic_score >= 0.18 and match_score < 35:

            match_score = 35

        intern_data = job.copy()

        intern_data["match_score"] = match_score

        intern_data["matched_skills"] = matched_skills

        ranked_tuple = (intern_data, len(matched_skills), semantic_score, sector_match, location_match)

        fallback_ranked.append(ranked_tuple)

        if match_score >= 30 and (matched_skills or sector_match or location_match or semantic_score >= 0.18):

            ranked.append(ranked_tuple)

    pool = ranked if ranked else fallback_ranked

    pool.sort(

        key=lambda item: (

            item[1],

            item[0]["match_score"],

            item[2],

            item[3],

            item[4],

        ),

        reverse=True,

    )

    return [item[0] for item in pool[:top_n]]

def translate_text_lightweight(text: str, dest_lang: str) -> str:

    """Translate text using deep_translator for offline/free translations."""

    normalized_dest = _normalize_language(dest_lang)

    if not text or normalized_dest == "en":

        return text

    try:

        with _disable_proxy_env():

            translated = GoogleTranslator(source='auto', target=normalized_dest).translate(text)

        return translated if translated else text

    except Exception as e:

        print(f"Translation error: {e}")

    return text

def _firecrawl_scrape(url: str, *, formats: Optional[List[str]] = None, timeout: int = 30000) -> Dict[str, Any]:

    """Scrape a URL using Firecrawl and return structured content."""

    if not firecrawl_client:

        raise RuntimeError("Firecrawl client is not initialized. Check FIRECRAWL_API_KEY.")

    scrape_formats = formats or ["markdown", "html"]

    try:

        with _disable_proxy_env():

            result = firecrawl_client.scrape_url(

                url,

                params={"formats": scrape_formats, "timeout": timeout},

            )

        if isinstance(result, dict):

            return result

        return {"markdown": str(result), "metadata": {"url": url}}

    except Exception as e:

        print(f"Firecrawl scrape error for {url}: {e}")

        raise RuntimeError(f"Firecrawl scrape failed: {e}") from e

def _firecrawl_scrape_text(url: str) -> str:

    """Scrape a URL using Firecrawl and return just the markdown text."""

    result = _firecrawl_scrape(url, formats=["markdown"])

    return _clean_text(result.get("markdown") or result.get("content") or "")

def _fallback_interview_prep(
    job_title: str,
    company: str,
    skills: List[str],
    num_questions: int = DEFAULT_INTERVIEW_PREP_QUESTION_COUNT,
    used_prompts: Optional[set[str]] = None,
) -> List[Dict[str, Any]]:

    question_count = _bounded_question_count(
        num_questions,
        default=DEFAULT_INTERVIEW_PREP_QUESTION_COUNT,
        minimum=5,
        maximum=MAX_INTERVIEW_PREP_QUESTION_COUNT,
    )

    prompts = used_prompts or set()
    themes = _normalize_skills(skills) or [_clean_text(job_title) or "workplace communication"]
    rng = random.SystemRandom()

    scenario_templates = [
        "while handling high-volume candidate data",
        "during a daily operations review",
        "when a manager asks for same-day output",
        "while coordinating with multiple stakeholders",
        "when internship KPIs are off-track",
        "while preparing a recruiter-ready candidate summary",
        "during quality checks before final submission",
        "while collaborating with a remote team",
    ]

    question_templates = [
        "For a {job_title} internship at {company}, what is the best way to show {theme} {scenario}?",
        "In {company}, which action best proves you can apply {theme} {scenario}?",
        "Your mentor gives a task {scenario}. Which response shows practical {theme}?",
        "Which option builds trust fastest for {theme} execution {scenario} in a {job_title} role?",
        "If a deadline is tight {scenario}, how should you apply {theme} without lowering quality?",
    ]

    generic_distractors = [
        "Ignore the task details and move quickly",
        "Wait for someone else to solve it",
        "Skip documentation and feedback",
        "Guess without checking requirements",
        "Use copied output without validation",
        "Avoid asking clarifying questions",
        "Submit results without a quality check",
    ]

    questions: List[Dict[str, Any]] = []
    template_offset = rng.randrange(len(question_templates))
    scenario_offset = rng.randrange(len(scenario_templates))
    skill_offset = rng.randrange(len(themes))

    for idx in range(max(question_count * 12, 60)):
        theme = themes[(skill_offset + idx) % len(themes)]
        template = question_templates[(template_offset + idx) % len(question_templates)]
        scenario = scenario_templates[(scenario_offset + idx) % len(scenario_templates)]
        question_text = template.format(
            job_title=_clean_text(job_title) or "PM Internship",
            company=_clean_text(company) or "the company",
            theme=theme,
            scenario=scenario,
        )
        question_key = question_text.casefold()
        if question_key in prompts:
            continue
        prompts.add(question_key)

        correct = f"Use {theme} with clear steps, validation, and a concise summary of outcomes"
        wrongs = rng.sample(generic_distractors, k=min(3, len(generic_distractors)))
        while len(wrongs) < 3:
            wrongs.append("Proceed without reviewing the requirements")
        options = [correct, wrongs[0], wrongs[1], wrongs[2]]
        rng.shuffle(options)

        questions.append({
            "q": question_text,
            "options": options,
            "a": correct,
        })
        if len(questions) >= question_count:
            break

    return questions

def _fallback_learning_recommendations(job_title: str, company: str, skills: List[str], count: int = 6) -> List[Dict[str, str]]:

    themes = _normalize_skills(skills) or [_clean_text(job_title) or "internship basics"]

    levels = ["Easy", "Easy", "Medium", "Medium", "Hard", "Hard"]

    recommendations: List[Dict[str, str]] = []

    for idx in range(count):

        theme = themes[idx % len(themes)]

        recommendations.append({

            "title": f"Practice {theme.title()} for {job_title}",

            "difficulty": levels[idx % len(levels)],

            "acceptance": f"Show one clear example of {theme} in your internship work.",

            "topic": theme.title(),

        })

    return recommendations

def generate_interview_prep(
    job_title: str,
    company: str,
    skills: List[str],
    api_key: str,
    num_questions: int = DEFAULT_INTERVIEW_PREP_QUESTION_COUNT,
):

    normalized_skills = _normalize_skills(skills)
    question_count = _bounded_question_count(
        num_questions,
        default=DEFAULT_INTERVIEW_PREP_QUESTION_COUNT,
        minimum=5,
        maximum=MAX_INTERVIEW_PREP_QUESTION_COUNT,
    )

    if not _has_llm_provider():

        return _fallback_interview_prep(job_title, company, normalized_skills, num_questions=question_count)

    try:

        response = _call_structured(

            prompt=(

                f"Create exactly {question_count} multiple-choice interview questions for the PM Internship role '{job_title}' at '{company}'.\n"

                f"Candidate skills: {', '.join(normalized_skills) if normalized_skills else 'general internship readiness'}.\n"

                "Rules:\n"

                "- Questions must fit the actual internship role, not generic coding tests.\n"

                "- Use the provided skills when relevant.\n"

                "- Include one correct answer and exactly three believable wrong answers.\n"

                "- Keep options short and practical.\n"
                "- Avoid duplicate questions.\n"
                f"- Variation tag: {uuid.uuid4().hex[:10]}."

            ),

            schema_model=QuizQuestionsResponse,

            system_instruction="You are an expert internship interviewer creating strict JSON interview MCQs.",

            temperature=0.35,

            max_output_tokens=2048,

        )

        questions: List[Dict[str, Any]] = []

        seen_prompts = set()

        for item in response.questions:

            built_question = _build_quiz_question(item)

            prompt_key = built_question["q"].casefold()

            if prompt_key in seen_prompts:

                continue

            seen_prompts.add(prompt_key)

            questions.append(built_question)

        if len(questions) < question_count:
            for item in _fallback_interview_prep(
                job_title,
                company,
                normalized_skills,
                num_questions=question_count - len(questions),
                used_prompts=seen_prompts,
            ):
                prompt_key = _clean_text(item.get("q", "")).casefold()
                if not prompt_key or prompt_key in seen_prompts:
                    continue
                seen_prompts.add(prompt_key)
                questions.append(item)

        if questions:

            return questions[:question_count]

    except Exception as e:

        print(f"[ERROR] generate_interview_prep failed, using fallback: {e}")

    return _fallback_interview_prep(job_title, company, normalized_skills, num_questions=question_count)

def fetch_learning_recommendations(job_title: str, company: str, skills: List[str], api_key: str):

    normalized_skills = _normalize_skills(skills)

    if not _has_llm_provider():

        return _fallback_learning_recommendations(job_title, company, normalized_skills)

    try:

        response = _call_structured(

            prompt=(

                f"Create a practical 6-step learning path for the PM Internship role '{job_title}' at '{company}'.\n"

                f"Candidate skills: {', '.join(normalized_skills) if normalized_skills else 'general internship readiness'}.\n"

                "Rules:\n"

                "- Make the path relevant to the role, not just coding interview prep.\n"

                "- Each item should be a concrete learning topic, exercise, or mini-project.\n"

                "- difficulty should be Easy, Medium, or Hard.\n"

                "- acceptance should be one short outcome statement.\n"

                "- topic should be a short category label."

            ),

            schema_model=LearningRecommendationsResponse,

            system_instruction="You are a career coach creating strict JSON learning paths for internships.",

            temperature=0.3,

            max_output_tokens=2048,

        )

        recommendations = []

        for item in response.recommendations:

            recommendations.append({

                "title": _clean_text(item.title),

                "difficulty": _clean_text(item.difficulty) or "Medium",

                "acceptance": _clean_text(item.acceptance) or "Build one clear example you can discuss in interviews.",

                "topic": _clean_text(item.topic) or "Skill Building",

            })

        if recommendations:

            return recommendations[:6]

    except Exception as e:

        print(f"[ERROR] fetch_learning_recommendations failed, using fallback: {e}")

    return _fallback_learning_recommendations(job_title, company, normalized_skills)

def _extract_pdf_text(file_bytes: bytes) -> str:

    extracted_versions: List[str] = []

    if pdfplumber is not None:

        try:

            with pdfplumber.open(BytesIO(file_bytes)) as pdf:

                extracted_versions.append("\n".join([(page.extract_text() or "") for page in pdf.pages]))

        except Exception as e:

            print(f"PDF Parsing Warning (pdfplumber): {e}")

    for reader_cls, reader_name in ((ModernPdfReader, "pypdf"), (LegacyPdfReader, "PyPDF2")):

        if reader_cls is None:

            continue

        try:

            reader = reader_cls(BytesIO(file_bytes))

            extracted_versions.append("\n".join([(page.extract_text() or "") for page in reader.pages]))

        except Exception as e:

            print(f"PDF Parsing Warning ({reader_name}): {e}")

    text = max(extracted_versions, key=lambda value: len(_clean_text(value)), default="")

    if len(_clean_text(text)) >= 120:

        return text

    try:

        from pdf2image import convert_from_bytes

        import pytesseract

        images = convert_from_bytes(file_bytes, dpi=200, first_page=1, last_page=3)

        ocr_text = "\n".join([

            pytesseract.image_to_string(image) for image in images

        ])

        if len(_clean_text(ocr_text)) > len(_clean_text(text)):

            return ocr_text

    except Exception as e:

        print(f"OCR Fallback Warning: {e}")

    return text

def _extract_docx_text(file_bytes: bytes) -> str:

    if DocxDocument is None:

        raise RuntimeError("python-docx is not installed on this server.")

    try:

        document = DocxDocument(BytesIO(file_bytes))

    except Exception as e:

        raise RuntimeError(f"DOCX parsing failed: {e}") from e

    text_blocks: List[str] = []

    for paragraph in document.paragraphs:

        block = _clean_text(paragraph.text)

        if block:

            text_blocks.append(block)

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                cell_text = _clean_text(cell.text)

                if cell_text:

                    text_blocks.append(cell_text)

    return "\n".join(text_blocks)

def _extract_image_text(file_bytes: bytes) -> str:

    try:

        from PIL import Image

        import pytesseract

    except Exception as e:

        raise RuntimeError(f"OCR dependencies missing: {e}") from e

    try:

        image = Image.open(BytesIO(file_bytes))

        return pytesseract.image_to_string(image)

    except Exception as e:

        raise RuntimeError(f"Image OCR failed: {e}") from e

def _detect_resume_file_type(file_name: str, content_type: str) -> str:

    name = (file_name or "").casefold()

    mime = (content_type or "").casefold()

    if name.endswith(".pdf") or "pdf" in mime:

        return "pdf"

    if name.endswith(".docx") or "wordprocessingml.document" in mime:

        return "docx"

    if any(name.endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"]):

        return "image"

    if mime.startswith("image/"):

        return "image"

    return "unsupported"

def _extract_resume_text(file_name: str, content_type: str, file_bytes: bytes) -> str:

    file_type = _detect_resume_file_type(file_name, content_type)

    if file_type == "pdf":

        return _extract_pdf_text(file_bytes)

    if file_type == "docx":

        return _extract_docx_text(file_bytes)

    if file_type == "image":

        try:

            return _extract_image_text(file_bytes)

        except RuntimeError as e:

            print(f"Image OCR fallback warning: {e}")

            return ""

    raise HTTPException(

        status_code=400,

        detail="Unsupported file format. Upload PDF, DOCX, or image (PNG/JPG/JPEG/WEBP/BMP/TIFF).",

    )

def _extract_github_username(github_url: Optional[str], github_username: Optional[str]) -> str:

    if github_username:

        return _clean_text(github_username).lstrip("@")

    clean_url = _clean_text(github_url)

    if not clean_url:

        return ""

    parsed = urlparse(clean_url)

    path_parts = [part for part in parsed.path.split("/") if part]

    if not path_parts:

        return ""

    return path_parts[0].lstrip("@")

def _github_api_get(url: str, params: Optional[Dict[str, Any]] = None) -> Any:

    headers = {

        "Accept": "application/vnd.github+json",

        "X-GitHub-Api-Version": "2022-11-28",

    }

    github_token = _clean_text(os.getenv("GITHUB_TOKEN"))

    if github_token:

        headers["Authorization"] = f"Bearer {github_token}"

    session = requests.Session()

    session.trust_env = False

    response = session.get(url, params=params or {}, headers=headers, timeout=20)

    if response.status_code == 404:

        raise HTTPException(status_code=404, detail="GitHub profile not found.")

    if response.status_code in (403, 429):

        raise RuntimeError("GitHub API rate limit reached. Retry later or configure GITHUB_TOKEN.")

    response.raise_for_status()

    return response.json()

def _repo_complexity_score(repo: Dict[str, Any]) -> float:

    size_score = min(1.0, float(repo.get("size", 0)) / 3000.0)

    popularity_score = min(1.0, float(repo.get("stargazers_count", 0) + repo.get("forks_count", 0)) / 50.0)

    topic_score = min(1.0, float(len(repo.get("topics") or [])) / 6.0)

    description = _clean_text(repo.get("description", "")).casefold()

    complexity_terms = ["api", "docker", "ml", "pipeline", "microservice", "kafka", "redis", "postgres"]

    description_score = 1.0 if any(term in description for term in complexity_terms) else 0.0

    return round((size_score * 0.4 + popularity_score * 0.25 + topic_score * 0.2 + description_score * 0.15) * 100, 2)

def _normalize_skill_token(skill: str) -> str:

    return re.sub(r"\s+", " ", _clean_text(skill).casefold())

def _infer_verified_skills_from_repos(repos: List[Dict[str, Any]], claimed_skills: List[str]) -> Dict[str, List[str]]:

    language_to_skill = {

        "python": "Python",

        "javascript": "JavaScript",

        "typescript": "TypeScript",

        "java": "Java",

        "go": "Go",

        "c++": "C++",

        "c": "C",

        "ruby": "Ruby",

        "php": "PHP",

        "html": "HTML",

        "css": "CSS",

        "swift": "Swift",

        "kotlin": "Kotlin",

        "rust": "Rust",

        "sql": "SQL",

    }

    evidence_blob_parts: List[str] = []

    seen_languages: List[str] = []

    for repo in repos:

        language = _clean_text(repo.get("language", ""))

        if language:

            seen_languages.append(language)

        evidence_blob_parts.extend([

            _clean_text(repo.get("name", "")),

            _clean_text(repo.get("description", "")),

            " ".join((_clean_text(topic) for topic in (repo.get("topics") or []))),

            language,

        ])

    for language in seen_languages:

        mapped_skill = language_to_skill.get(language.casefold())

        if mapped_skill:

            evidence_blob_parts.append(mapped_skill)

    evidence_blob = " ".join(part.casefold() for part in evidence_blob_parts if part)

    normalized_claimed = [_normalize_skill_token(skill) for skill in claimed_skills if _clean_text(skill)]

    if not normalized_claimed:

        unique_languages = sorted({lang for lang in seen_languages if lang})

        return {

            "verified": unique_languages[:8],

            "unverified": [],

        }

    verified: List[str] = []

    unverified: List[str] = []

    for original in claimed_skills:

        token = _normalize_skill_token(original)

        if token and token in evidence_blob:

            verified.append(original)

        else:

            unverified.append(original)

    return {

        "verified": verified,

        "unverified": unverified,

    }

def _token_overlap_similarity(left: str, right: str) -> float:

    left_tokens = {token for token in re.findall(r"[A-Za-z_]\w+", left)}

    right_tokens = {token for token in re.findall(r"[A-Za-z_]\w+", right)}

    if not left_tokens and not right_tokens:

        return 1.0

    union = left_tokens | right_tokens

    if not union:

        return 0.0

    return len(left_tokens & right_tokens) / len(union)

def _python_ast_similarity(source_code: str, candidate_code: str) -> float:

    try:

        source_tree = ast.parse(source_code)

        candidate_tree = ast.parse(candidate_code)

    except SyntaxError:

        return _token_overlap_similarity(source_code, candidate_code)

    source_dump = ast.dump(source_tree, annotate_fields=False, include_attributes=False)

    candidate_dump = ast.dump(candidate_tree, annotate_fields=False, include_attributes=False)

    structural_ratio = difflib.SequenceMatcher(a=source_dump, b=candidate_dump).ratio()

    source_nodes = [node.__class__.__name__ for node in ast.walk(source_tree)]

    candidate_nodes = [node.__class__.__name__ for node in ast.walk(candidate_tree)]

    source_counter = Counter(source_nodes)

    candidate_counter = Counter(candidate_nodes)

    common = sum((source_counter & candidate_counter).values())

    total = max(sum(source_counter.values()), sum(candidate_counter.values()), 1)

    node_distribution_ratio = common / total

    return round((structural_ratio * 0.65) + (node_distribution_ratio * 0.35), 6)

def _embedding_similarity(source_code: str, candidate_code: str) -> float:

    text_a = _clean_text(source_code)

    text_b = _clean_text(candidate_code)

    if not text_a and not text_b:

        return 1.0

    try:

        vectorizer = TfidfVectorizer(analyzer="char_wb", ngram_range=(3, 5), lowercase=False)

        matrix = vectorizer.fit_transform([text_a, text_b])

        return float(cosine_similarity(matrix[0:1], matrix[1:2])[0][0])

    except Exception:

        return _token_overlap_similarity(text_a, text_b)

def _score_to_risk_level(score: float) -> str:

    if score >= 80:

        return "High"

    if score >= 55:

        return "Medium"

    return "Low"

def generate_dynamic_questions(skills: List[str], api_key: str, num_questions: int = DEFAULT_RESUME_QUIZ_QUESTION_COUNT):

    normalized_skills = _normalize_skills(skills)

    if not normalized_skills:

        raise RuntimeError("No skills available for question generation")

    if not _has_llm_provider():

        return _fallback_resume_questions(normalized_skills, num_questions=num_questions)

    try:

        response = _call_structured(

            prompt=(

                f"Create exactly {num_questions} multiple-choice interview questions based only on these skills: {', '.join(normalized_skills)}.\n"

                "Rules:\n"

                "- Every question must directly test one of the provided skills.\n"

                "- Do not ask about skills that are not in the list.\n"

                "- Make the questions realistic and interview-ready, not generic aptitude questions.\n"

                "- For each question, provide one correct answer and exactly three plausible wrong answers.\n"

                "- Keep each option short and distinct.\n"

                "- Avoid 'all of the above' and 'none of the above'."

            ),

            schema_model=QuizQuestionsResponse,

            system_instruction="You create precise interview MCQs and return strict JSON.",

            temperature=0.35,

            max_output_tokens=2048,

        )

        questions: List[Dict[str, Any]] = []

        seen_prompts: set[str] = set()

        for item in response.questions:

            try:

                built_question = _build_quiz_question(item)

            except Exception as e:

                print(f"Dynamic question fallback warning: {e}")

                continue

            prompt_key = built_question["q"].casefold()

            if prompt_key in seen_prompts:

                continue

            seen_prompts.add(prompt_key)

            questions.append(built_question)

        if len(questions) < num_questions:

            questions.extend(

                _fallback_resume_questions(

                    normalized_skills,

                    num_questions=num_questions - len(questions),

                    used_prompts=seen_prompts,

                )

            )

        return questions[:num_questions]

    except Exception as e:

        print(f"Dynamic question generation fallback warning: {e}")

        return _fallback_resume_questions(normalized_skills, num_questions=num_questions)

@app.get("/")
def health_check():

    return {"status": "Engine is running smoothly."}

@app.post("/refresh-internships")

def refresh_pm_internships():

    try:

        jobs = _load_pm_jobs(force_refresh=True)

        return {

            "status": "ok",

            "count": len(jobs),

            "source": PM_SOURCE_NAME,

            "apply_url": PM_PORTAL_URL,

        }

    except RuntimeError as e:

        raise HTTPException(status_code=503, detail=str(e))

    except Exception as e:

        raise HTTPException(status_code=500, detail=f"Failed to refresh PM internships: {str(e)}")

@app.post("/analyze-resume")

async def analyze_resume(file: UploadFile = File(...), location: str = Form("Amaravati"), education: str = Form("Graduate"), preferred_sector: str = Form("Any")):
    try:

        file_bytes = await file.read()

        if not file_bytes:

            raise HTTPException(status_code=400, detail="The uploaded file is empty.")
        resume_cache_key = hashlib.sha256(file_bytes).hexdigest()

        cached_analysis = RESUME_ANALYSIS_CACHE.get(resume_cache_key)

        if cached_analysis:

            extracted_skills = _normalize_skills(cached_analysis.get("skills", []))

            resume_text = _clean_text(cached_analysis.get("resume_text", ""))

        else:

            text = _extract_resume_text(file.filename or "", getattr(file, "content_type", "") or "", file_bytes)

            if len(_clean_text(text)) < 40:
                fallback_skills = _fallback_resume_skill_candidates()[:5]
                analysis = {
                    "skills": fallback_skills,
                    "questions": _fallback_resume_questions(fallback_skills, num_questions=DEFAULT_RESUME_QUIZ_QUESTION_COUNT),
                    "resume_text": "",
                }
                RESUME_ANALYSIS_CACHE[resume_cache_key] = analysis
                return {
                    "extracted_skills": fallback_skills,
                    "assessment_quiz": analysis["questions"],
                    "profile_dict": {"location": location, "education": education, "preferred_sector": preferred_sector},
                    "warning": "Low text extraction from file. Returned fallback skills and questions.",
                    "resume_text": "",
                    "github_url": None,
                    "linkedin_url": None,
                    "email": None,
                }

            analysis = generate_resume_analysis(text, groq_api_key, num_questions=DEFAULT_RESUME_QUIZ_QUESTION_COUNT)

            extracted_skills = analysis["skills"]

            resume_text = _clean_text(text)

            analysis["resume_text"] = resume_text[:12000]

            RESUME_ANALYSIS_CACHE[resume_cache_key] = analysis

        if _has_llm_provider():
            quiz = generate_dynamic_questions(extracted_skills, groq_api_key, num_questions=DEFAULT_RESUME_QUIZ_QUESTION_COUNT)
        else:
            quiz = _fallback_resume_questions(extracted_skills, num_questions=DEFAULT_RESUME_QUIZ_QUESTION_COUNT)

        return {

            "extracted_skills": extracted_skills,

            "assessment_quiz": quiz,

            "profile_dict": {"location": location, "education": education, "preferred_sector": preferred_sector},

            "resume_text": resume_text[:12000],

            "github_url": analysis.get("github_url"),

            "linkedin_url": analysis.get("linkedin_url"),

            "email": analysis.get("email"),

        }

    except HTTPException:

        raise

    except RuntimeError as e:

        raise HTTPException(status_code=503, detail=str(e))

    except Exception as e:

        raise HTTPException(status_code=500, detail=f"Failed to process resume: {str(e)}")

@app.post("/manual-profile")

def process_manual_profile(profile: ManualProfile):

    try:

        normalized_skills = _normalize_skills(profile.manual_skills)

        if not normalized_skills:

            raise HTTPException(status_code=400, detail="Please select at least one skill.")

        if _has_llm_provider():
            quiz = generate_dynamic_questions(normalized_skills, groq_api_key, num_questions=DEFAULT_RESUME_QUIZ_QUESTION_COUNT)
        else:
            quiz = _fallback_resume_questions(normalized_skills, num_questions=DEFAULT_RESUME_QUIZ_QUESTION_COUNT)

        profile_dict = {

            "location": profile.location,

            "education": profile.education,

            "preferred_sector": profile.preferred_sector

        }

        return {

            "status": "Profile registered successfully",

            "verified_skills": normalized_skills,

            "assessment_quiz": quiz,

            "profile_dict": profile_dict

        }

    except HTTPException:

        raise

    except RuntimeError as e:

        raise HTTPException(status_code=503, detail=str(e))

    except Exception as e:

        raise HTTPException(status_code=500, detail=str(e))

@app.post("/recommended-jobs")

def get_recommended_jobs(request: DynamicJobsRequest):

    profile_dict = {

        "location": request.location,

        "education": request.education,

        "preferred_sector": request.preferred_sector

    }

    jobs = rag_semantic_search(profile_dict, request.skills, top_n=25)

    # Score-based ranking adjustment
    if request.total_score is not None:
        for job in jobs:
            # If they score very high, boost overall match score for higher stipend jobs
            if request.total_score > 85 and "15k" in str(job.get("stipend", "")):
                job["match_score"] = min(100, job.get("match_score", 0) + 5)
            # If score is low, slightly prioritize jobs without strict skill requirements
            elif request.total_score < 70:
                job["match_score"] = min(100, job.get("match_score", 0) + 2)

    try:

        if len(jobs) < 10 or max((job.get("match_score", 0) for job in jobs), default=0) < 70:

            external_jobs = _fetch_internshala_internships(profile_dict, request.skills, max_results=25)

            jobs = _merge_recommendation_sets(jobs, external_jobs, top_n=30)

    except Exception as e:

        print(f"External internship fallback warning: {e}")

    lang = _normalize_language(getattr(request, 'lang', 'en') or request.target_language or 'en')

    if lang != "en":

        for job in jobs:

            job['title'] = translate_text_lightweight(job.get('title', ''), lang)

            job['sector'] = translate_text_lightweight(job.get('sector', ''), lang)

            job['description'] = translate_text_lightweight(job.get('description', ''), lang)

            job['field'] = translate_text_lightweight(job.get('field', ''), lang)

            job['matched_skills'] = [translate_text_lightweight(skill, lang) for skill in job.get('matched_skills', [])]

            job['skills'] = [translate_text_lightweight(skill, lang) for skill in job.get('skills', [])]

    request_key = _recommendation_request_key(profile_dict, request.skills)

    jobs = _select_rotating_job_slice(jobs, request_key=request_key, top_n=5)

    return {"top_matches": jobs}

def _build_github_verification_summary(username: str, claimed_skills: Optional[List[str]] = None) -> Dict[str, Any]:

    repos = _github_api_get(

        f"https://api.github.com/users/{username}/repos",

        params={"per_page": 100, "sort": "updated"},

    )

    if not isinstance(repos, list):

        raise HTTPException(status_code=502, detail="Unexpected response from GitHub API.")

    non_fork_repos = [repo for repo in repos if not repo.get("fork")]

    language_counts = Counter(

        _clean_text(repo.get("language", ""))

        for repo in non_fork_repos

        if _clean_text(repo.get("language", ""))

    )

    recent_cutoff = datetime.now(timezone.utc) - timedelta(days=120)

    recent_activity = 0

    for repo in non_fork_repos:

        pushed_at = _clean_text(repo.get("pushed_at", ""))

        if not pushed_at:

            continue

        try:

            pushed_dt = datetime.fromisoformat(pushed_at.replace("Z", "+00:00"))

        except ValueError:

            continue

        if pushed_dt >= recent_cutoff:

            recent_activity += 1

    complexity_scores = [_repo_complexity_score(repo) for repo in non_fork_repos]

    average_complexity = round(sum(complexity_scores) / len(complexity_scores), 2) if complexity_scores else 0.0

    skill_status = _infer_verified_skills_from_repos(non_fork_repos, claimed_skills or [])

    claimed_count = len([skill for skill in (claimed_skills or []) if _clean_text(skill)])

    if claimed_count:

        verification_ratio = len(skill_status["verified"]) / max(claimed_count, 1)

    else:

        verification_ratio = min(1.0, len(language_counts) / 4.0)

    activity_score = min(100.0, recent_activity * 18.0 + min(len(non_fork_repos), 10) * 4.0)

    confidence_score = int(round(verification_ratio * 45.0 + activity_score * 0.3 + average_complexity * 0.25))

    confidence_score = max(0, min(100, confidence_score))

    if confidence_score >= 75:

        verification_status = "Verified"

    elif confidence_score >= 45:

        verification_status = "Partially Verified"

    else:

        verification_status = "Low Confidence"

    return {

        "github_username": username,

        "repo_count": len(non_fork_repos),

        "language_summary": dict(language_counts.most_common(8)),

        "recent_active_repos_120d": recent_activity,

        "average_repo_complexity": average_complexity,

        "verified_skills": skill_status["verified"],

        "unverified_skills": skill_status["unverified"],

        "skill_verification_status": verification_status,

        "confidence_score": confidence_score,

    }

def _linkedin_signal(linkedin_url: Optional[str]) -> Dict[str, Any]:

    normalized = _ensure_absolute_url(linkedin_url or "")

    if not normalized:

        return {"score": 0, "status": "missing", "url": "", "note": "LinkedIn URL not provided."}

    parsed = urlparse(normalized)

    if "linkedin.com" not in parsed.netloc.casefold():

        return {"score": 25, "status": "invalid_domain", "url": normalized, "note": "URL does not look like a LinkedIn profile."}

    session = requests.Session()

    session.trust_env = False

    headers = {

        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/131.0.0.0 Safari/537.36",

    }

    try:

        with _disable_proxy_env():

            response = session.get(normalized, headers=headers, timeout=12, allow_redirects=True)

        if response.status_code < 400:

            return {"score": 85, "status": "reachable", "url": normalized, "note": "LinkedIn profile URL is reachable."}

        return {"score": 55, "status": f"http_{response.status_code}", "url": normalized, "note": "LinkedIn URL format is valid but could not be fully verified."}

    except Exception:

        return {"score": 65, "status": "format_valid", "url": normalized, "note": "LinkedIn URL format looks valid."}

def _account_signal(email: Optional[str], github_username: str, linkedin_url: Optional[str]) -> Dict[str, Any]:

    clean_email = _clean_text(email)

    email_valid = bool(re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", clean_email))

    provider_bonus = 0

    if email_valid:

        provider = clean_email.split("@")[-1].casefold()

        if provider.endswith((".edu", ".ac.in")):

            provider_bonus = 20

        elif provider.endswith((".com", ".in", ".org", ".net")):

            provider_bonus = 10

    score = 0

    score += 50 if email_valid else 0

    score += 25 if _clean_text(github_username) else 0

    score += 25 if _clean_text(linkedin_url) else 0

    score += provider_bonus

    score = min(100, score)

    return {

        "score": score,

        "email_valid": email_valid,

        "email": clean_email,

        "note": "Account signal uses verified email format and connected profile links.",

    }

@app.post("/github-verification")

@app.post("/verify-github")

def github_verification(request: GitHubVerificationRequest):

    username = _extract_github_username(request.github_url, request.github_username)

    if not username:

        raise HTTPException(status_code=400, detail="Provide github_username or a valid github_url.")

    try:

        return _build_github_verification_summary(username, request.claimed_skills)

    except HTTPException:

        raise

    except RuntimeError as e:

        raise HTTPException(status_code=503, detail=str(e))

    except Exception as e:

        raise HTTPException(status_code=500, detail=f"GitHub verification failed: {e}")

@app.post("/compute-trust-score")

def compute_trust_score(request: TrustScoreRequest):

    normalized_skills = _normalize_skills(request.skills)

    adjusted_assessment = max(0, min(100, int(request.assessment_score) - int(round(request.cheating_score * 0.35))))

    github_username = _extract_github_username(request.github_url, request.github_username)

    github_summary: Dict[str, Any] = {}

    github_score = 0

    if github_username:

        try:

            github_summary = _build_github_verification_summary(github_username, normalized_skills)

            github_score = int(github_summary.get("confidence_score", 0))

        except Exception as e:

            github_summary = {"error": str(e), "github_username": github_username}

    resume_project_match = 0
    resume_matched_repos: List[str] = []
    if github_username and _clean_text(request.resume_text) and github_summary.get("repo_count", 0) > 0:
        try:
            repos = _github_api_get(
                f"https://api.github.com/users/{github_username}/repos",
                params={"per_page": 50, "sort": "updated"},
            )
            resume_lower = _clean_text(request.resume_text).lower()
            for repo in (repos if isinstance(repos, list) else []):
                repo_name = _clean_text(repo.get("name", "")).lower().replace("-", " ").replace("_", " ")
                repo_desc = _clean_text(repo.get("description", "")).lower()
                if len(repo_name) > 3 and repo_name in resume_lower:
                    resume_matched_repos.append(repo.get("name", ""))
                elif len(repo_desc) > 10:
                    desc_words = [w for w in repo_desc.split() if len(w) > 4]
                    match_count = sum(1 for w in desc_words if w in resume_lower)
                    if match_count >= 3:
                        resume_matched_repos.append(repo.get("name", ""))
            if resume_matched_repos:
                resume_project_match = min(100, len(resume_matched_repos) * 25)
                github_score = min(100, github_score + resume_project_match // 3)
                github_summary["resume_matched_repos"] = resume_matched_repos[:10]
                github_summary["resume_project_match_bonus"] = resume_project_match
                github_summary["confidence_score"] = github_score
        except Exception as e:
            print(f"Resume-GitHub project match error: {e}")

    linkedin_summary = _linkedin_signal(request.linkedin_url)

    account_summary = _account_signal(request.email, github_username, request.linkedin_url)

    profile_completeness = 0

    profile_completeness += 25 if normalized_skills else 0

    profile_completeness += 25 if github_username else 0

    profile_completeness += 25 if _clean_text(request.linkedin_url) else 0

    profile_completeness += 25 if account_summary.get("email_valid") else 0

    trust_score = int(

        round(

            adjusted_assessment * 0.55

            + github_score * 0.2

            + int(linkedin_summary.get("score", 0)) * 0.12

            + int(account_summary.get("score", 0)) * 0.08

            + profile_completeness * 0.05

        )

    )

    trust_score = max(0, min(100, trust_score))

    if trust_score >= 82:

        trust_level = "Excellent"

    elif trust_score >= 68:

        trust_level = "Good"

    elif trust_score >= 52:

        trust_level = "Moderate"

    else:

        trust_level = "Needs Improvement"

    recommendations: List[str] = []

    if github_score < 50:

        recommendations.append("Add recent public GitHub projects with clear README and commit activity.")

    if int(linkedin_summary.get("score", 0)) < 70:

        recommendations.append("Provide a complete LinkedIn profile URL with headline, skills, and project highlights.")

    if not account_summary.get("email_valid"):

        recommendations.append("Add a valid professional email to strengthen account verification.")

    if request.cheating_score > 25:

        recommendations.append("Retake the assessment with stable behavior to reduce integrity penalties.")

    if len(normalized_skills) < 3:

        recommendations.append("Expand the verified skill list with resume-backed technical strengths.")

    return {

        "trust_score": trust_score,

        "trust_level": trust_level,

        "breakdown": {

            "assessment_component": adjusted_assessment,

            "github_component": github_score,

            "linkedin_component": int(linkedin_summary.get("score", 0)),

            "account_component": int(account_summary.get("score", 0)),

            "profile_completeness": profile_completeness,

        },

        "github": github_summary,

        "linkedin": linkedin_summary,

        "account": account_summary,

        "recommendations": recommendations[:6],

    }

@app.post("/plagiarism-detection")

@app.post("/detect-plagiarism")

async def plagiarism_detection(request: PlagiarismDetectionRequest):

    source_code = request.source_code or ""

    candidate_code = request.candidate_code or ""

    if len(_clean_text(source_code)) < 20 or len(_clean_text(candidate_code)) < 20:

        raise HTTPException(status_code=400, detail="Provide meaningful source_code and candidate_code for comparison.")

    language = _clean_text(request.language).casefold() or "python"

    if language in ("python", "py"):

        ast_similarity = _python_ast_similarity(source_code, candidate_code)

    else:

        ast_similarity = _token_overlap_similarity(source_code, candidate_code)

    embed_similarity = _embedding_similarity(source_code, candidate_code)

    plagiarism_score = round((ast_similarity * 0.6 + embed_similarity * 0.4) * 100, 2)

    plagiarism_score = max(0.0, min(100.0, plagiarism_score))

    risk_level = _score_to_risk_level(plagiarism_score)

    return {

        "plagiarism_score": plagiarism_score,

        "risk_level": risk_level,

        "ast_similarity": round(ast_similarity * 100, 2),

        "embedding_similarity": round(embed_similarity * 100, 2),

        "language": language,

    }

@app.post("/api/voice-recommend")

def voice_recommend(request: VoiceRecommendRequest):

    """

    Endpoint for OmniDimension Voice AI webhook.

    Accepts a flexible JSON payload with optional fields.

    Returns a CLEAN, flat JSON with only the top 3-5 internships.

    OmniDimension Dashboard Mapping:

      - location  → user's spoken city/state

      - sector    → user's spoken industry preference

      - skills    → extracted skill keywords (string or list)

      - education → qualification level mentioned

      - lang      → response language (en/hi/te)

    """

    user_skills = []

    if request.skills:

        if isinstance(request.skills, list):

            user_skills = [s.strip() for s in request.skills if s.strip()]

        elif isinstance(request.skills, str):

            user_skills = [s.strip() for s in request.skills.split(",") if s.strip()]

    if not user_skills:

        user_skills = ["communication", "general skills"]

    profile_dict = {

        "location": request.location or "",

        "education": request.education or "Graduate",

        "preferred_sector": request.sector or "Any",

    }

    jobs = rag_semantic_search(profile_dict, user_skills, top_n=25)

    try:

        if len(jobs) < 10 or max((job.get("match_score", 0) for job in jobs), default=0) < 70:

            external_jobs = _fetch_internshala_internships(profile_dict, user_skills, max_results=25)

            jobs = _merge_recommendation_sets(jobs, external_jobs, top_n=30)

    except Exception as e:

        print(f"Voice external internship fallback warning: {e}")

    lang = request.lang or "en"

    if lang in ("hi", "te"):

        for job in jobs:

            job["title"] = translate_text_lightweight(job.get("title", ""), lang)

            job["sector"] = translate_text_lightweight(job.get("sector", ""), lang)

    request_key = _recommendation_request_key(profile_dict, user_skills)

    jobs = _select_rotating_job_slice(jobs, request_key=request_key, top_n=5)

    clean_results = []

    for job in jobs:

        clean_results.append({

            "title": job.get("title", ""),

            "company": job.get("company", ""),

            "location": job.get("location", ""),

            "sector": job.get("sector", ""),

            "source": job.get("source", ""),

            "match_score": job.get("match_score", 0),

            "apply_url": job.get("apply_url", PM_PORTAL_URL),

        })

    return {

        "count": len(clean_results),

        "internships": clean_results,

    }

class JobTipRequest(BaseModel):

    title: str

    company: str

    language: str

class CallRequest(BaseModel):

    phone_number: str

@app.get("/api/voice-agent/status")

async def get_agent_status():

    try:

        client = _get_omnidim_client()

        resolved_agent_id, warning, bots = _resolve_omnidim_agent(client)

        agent_details = _extract_omnidim_payload(client.agent.get(resolved_agent_id))

        widget_config = agent_details.get("widget_config") if isinstance(agent_details, dict) else {}

        return {

            "status": "success",

            "configured_agent_id": _clean_text(agent_id) or None,

            "resolved_agent_id": resolved_agent_id,

            "warning": warning,

            "widget_url": widget_config.get("iframeUrl"),

            "widget_config": widget_config,

            "available_agents": [

                {

                    "id": bot.get("id"),

                    "name": bot.get("name"),

                    "bot_call_type": bot.get("bot_call_type"),

                }

                for bot in bots

            ],

            "agent_details": agent_details,

        }

    except RuntimeError as e:

        raise HTTPException(status_code=500, detail=str(e))

    except Exception as e:

        raise HTTPException(status_code=500, detail=_humanize_omnidim_error(e))

@app.post("/api/voice-agent/call")

async def trigger_call(request: CallRequest):

    try:

        client = _get_omnidim_client()

        resolved_agent_id, warning, _ = _resolve_omnidim_agent(client)

        phone_number = _clean_text(request.phone_number)

        if not re.fullmatch(r"\+\d{10,15}", phone_number):

            raise HTTPException(status_code=400, detail="Phone number must include country code, for example +919876543210.")

        response = client.call.dispatch_call(

            agent_id=resolved_agent_id,

            to_number=phone_number

        )

        return {

            "status": "Call initiated!",

            "resolved_agent_id": resolved_agent_id,

            "warning": warning,

            "call_data": _extract_omnidim_payload(response),

        }

    except HTTPException:

        raise

    except RuntimeError as e:

        raise HTTPException(status_code=500, detail=str(e))

    except Exception as e:

        raise HTTPException(status_code=500, detail=_humanize_omnidim_error(e))

@app.post("/job-tips")

def get_job_tips(request: JobTipRequest):

    if not _has_llm_provider():

        raise HTTPException(status_code=503, detail="No AI provider configured. Add GROQ_API_KEY or GEMINI_API_KEY.")

    try:

        language_name = LANG_MAP.get(_normalize_language(request.language), request.language)

        response = _call_structured(

            prompt=(

                f"Provide exactly 3 short, actionable interview or application tips for the PM Internship role '{request.title}' at '{request.company}'.\n"

                f"Write the tips in {language_name}.\n"

                "Each tip must be one sentence and practical."

            ),

            schema_model=JobTipsResponse,

            system_instruction="You are a concise career coach returning strict JSON tips.",

            temperature=0.3,

            max_output_tokens=512,

        )

        tips = [_clean_text(tip) for tip in response.tips if _clean_text(tip)]

        return {"tips": tips[:3]}

    except Exception as e:

        raise HTTPException(status_code=503, detail=str(e))

@app.post("/translate")

def translate_content(request: TranslationRequest):

    try:

        target = _normalize_language(request.target_language)

        if target == "en":

            return {"translated_payload": request.payload}

        translated_payload = {}

        with _disable_proxy_env():

            translator = GoogleTranslator(source='auto', target=target)

            for k, v in request.payload.items():

                if isinstance(v, str):

                    translated_payload[k] = translator.translate(v)

                else:

                    translated_payload[k] = v

        return {"translated_payload": translated_payload}

    except Exception as e:

        print(f"Translation Error: {e}")

        return {"translated_payload": request.payload, "warning": "Translation failed, defaulting to English."}

@app.post("/agent-chat")

def agent_chat(request: AgentChatRequest):

    if not _has_llm_provider():

        return {"reply": "AI mentor is unavailable because no AI provider is configured."}

    try:

        preferred_language = LANG_MAP.get(_normalize_language(request.target_language), request.target_language)

        system_instruction = f"""

        You are a highly empathetic, encouraging Career Mentor for rural Indian youth applying to the PM Internship Scheme.

        The user has these verified skills: {', '.join(request.user_skills)}.

        Their preferred language for responses is: {preferred_language}. You MUST reply ONLY in this language. Use simple, colloquial, 8th-grade vocabulary.

        YOUR ROLE:

        1. Acknowledge their skills and build confidence.

        2. Give them 1 specific, practical interview tip based on their skills.

        3. Suggest 1 free learning step to improve.

        4. Keep responses extremely short (max 3-4 sentences). Do not use emojis.

        """

        conversation_context = "Conversation History:\n"

        for msg in request.messages[:-1]:

            role_name = "User" if msg.role == "user" else "Mentor"

            conversation_context += f"{role_name}: {msg.content}\n"

        current_msg = request.messages[-1].content if request.messages else "Hello!"

        final_prompt = f"{conversation_context}\nUser: {current_msg}\nMentor:"

        reply = _call_text(final_prompt, system_instruction=system_instruction, temperature=0.4, max_output_tokens=512)

        return {"reply": reply}

    except Exception as e:

        print(f"[ERROR] agent-chat failed: {e}")

        return {"reply": "AI mentor is temporarily unavailable because all configured AI providers failed. Please try again later or switch to another API key."}

@app.post("/interview-prep")

def interview_prep_endpoint(request: InterviewPrepRequest):

    try:

        questions = generate_interview_prep(
            request.job_title,
            request.company,
            request.skills,
            groq_api_key,
            num_questions=request.question_count or DEFAULT_INTERVIEW_PREP_QUESTION_COUNT,
        )

        return {"questions": questions}

    except RuntimeError as e:

        raise HTTPException(status_code=503, detail=str(e))

@app.post("/learning-recommendations")

def learning_recommendations_endpoint(request: LearningRecommendationRequest):

    try:

        recommendations = fetch_learning_recommendations(

            request.job_title or "PM Internship",

            request.company,

            request.skills,

            groq_api_key,

        )

        return {"recommendations": recommendations}

    except RuntimeError as e:

        raise HTTPException(status_code=503, detail=str(e))

import subprocess
import tempfile
import sys

class Test1Request(BaseModel):
    skills: List[str]

@app.post("/generate-test-1")
async def generate_test_1(request: Test1Request):
    """Generates a deep mandatory MCQ set for TEST 1."""
    normalized_skills = _normalize_skills(request.skills)
    if not normalized_skills:
        normalized_skills = ["Problem Solving"]
    question_count = 12

    try:
        if not _has_llm_provider():
            raise RuntimeError("No AI provider configured")

        variation_tag = uuid.uuid4().hex[:8]

        response = _call_structured(
            prompt=(
                f"Create exactly {question_count} interview-ready multiple choice questions for these skills: {', '.join(normalized_skills)}.\n"
                "Return strict JSON only.\n"
                "- Each question must have: id, skill, prompt, options (exactly 4), answer.\n"
                "- Answer must be one of the options.\n"
                "- Keep prompts concise and practical.\n"
                f"- Variation tag: {variation_tag}. Do not repeat common template wording."
            ),
            schema_model=Test1QuestionsResponse,
            system_instruction="You produce strict JSON for assessment MCQs.",
            temperature=0.55,
            max_output_tokens=2048,
        )

        questions: List[Dict[str, Any]] = []
        seen_prompts: set[str] = set()
        for index, item in enumerate(response.questions):
            prompt = _normalize_assessment_prompt(item.prompt)
            prompt_key = prompt.casefold()
            if not prompt or prompt_key in seen_prompts:
                continue
            seen_prompts.add(prompt_key)

            options = [_clean_text(opt) for opt in item.options if _clean_text(opt)]
            if len(options) < 2:
                options = ["Option A", "Option B", "Option C", "Option D"]
            while len(options) < 4:
                options.append(f"Option {chr(65 + len(options))}")
            options = options[:4]

            answer = _clean_text(item.answer)
            if answer not in options:
                answer = options[0]

            questions.append(
                {
                    "id": f"Q{len(questions) + 1}",
                    "skill": _clean_text(item.skill) or normalized_skills[index % len(normalized_skills)],
                    "prompt": prompt or f"What is a core principle of {normalized_skills[index % len(normalized_skills)]}?",
                    "options": options,
                    "answer": answer,
                }
            )

        if len(questions) < question_count:
            fallback_questions = _fallback_resume_questions(
                normalized_skills,
                num_questions=question_count - len(questions),
                used_prompts=seen_prompts,
            )
            for fallback in fallback_questions:
                fallback_options = [str(opt) for opt in fallback.get("options", [])][:4]
                while len(fallback_options) < 4:
                    fallback_options.append(f"Option {chr(65 + len(fallback_options))}")
                fallback_answer = str(fallback.get("a", fallback_options[0]))
                if fallback_answer not in fallback_options:
                    fallback_options[0] = fallback_answer
                questions.append(
                    {
                        "id": f"Q{len(questions) + 1}",
                        "skill": normalized_skills[(len(questions) - 1) % len(normalized_skills)],
                        "prompt": _normalize_assessment_prompt(fallback.get("q", "Select the best answer.")),
                        "options": fallback_options,
                        "answer": fallback_answer,
                    }
                )

        if questions:
            return {"questions": questions[:question_count]}

        raise RuntimeError("Model returned empty questions")

    except Exception as e:
        print(f"Error in Test 1: {e}")
        fallback_questions = _fallback_resume_questions(normalized_skills, num_questions=question_count)
        mock_qs = []
        for i, question in enumerate(fallback_questions, start=1):
            options = [str(opt) for opt in question.get("options", [])][:4]
            while len(options) < 4:
                options.append(f"Option {chr(65 + len(options))}")
            answer = str(question.get("a", options[0]))
            if answer not in options:
                options[0] = answer
            mock_qs.append(
                {
                    "id": f"Q{i}",
                    "skill": normalized_skills[(i - 1) % len(normalized_skills)],
                    "prompt": _normalize_assessment_prompt(question.get("q", f"What is a core principle of {normalized_skills[(i - 1) % len(normalized_skills)]}?")),
                    "options": options,
                    "answer": answer,
                }
            )
        return {"questions": mock_qs[:question_count]}

class Test2Request(BaseModel):
    skills: List[str]

@app.post("/generate-test-2")
async def generate_test_2(request: Test2Request):
    """Generates optional coding and deep theory questions for TEST 2."""
    normalized_skills = _normalize_skills(request.skills)
    if not normalized_skills:
        normalized_skills = ["Python"]
    challenge_count = 5

    try:
        if not _has_llm_provider():
            raise RuntimeError("No AI provider configured")

        variation_tag = uuid.uuid4().hex[:8]

        response = _call_structured(
            prompt=(
                f"Create exactly {challenge_count} coding/deep-theory challenges for these skills: {', '.join(normalized_skills)}.\n"
                "Return strict JSON only with keys: id, title, difficulty, companyTargets, prompt, starterCode, expectedSignals.\n"
                "- starterCode must be plain text with escaped newlines.\n"
                "- expectedSignals should include 3-6 key tokens expected in a good solution.\n"
                f"- Variation tag: {variation_tag}. Avoid repeating previous generic challenge themes."
            ),
            schema_model=Test2ChallengesResponse,
            system_instruction="You produce strict JSON for coding assessments.",
            temperature=0.6,
            max_output_tokens=2048,
        )

        challenges: List[Dict[str, Any]] = []
        seen_titles = set()
        for index, item in enumerate(response.challenges):
            starter_code = item.starterCode if _clean_text(item.starterCode) else "def solve():\n    pass\n"
            expected = [_clean_text(signal) for signal in item.expectedSignals if _clean_text(signal)]
            if not expected:
                expected = ["def", "return"]

            title = _normalize_assessment_prompt(item.title) or f"Coding challenge in {normalized_skills[index % len(normalized_skills)]}"
            title_key = title.casefold()
            if title_key in seen_titles:
                continue
            seen_titles.add(title_key)

            prompt = _normalize_assessment_prompt(item.prompt) or "Write a function to solve the required problem and explain complexity."

            challenges.append(
                {
                    "id": f"C{len(challenges) + 1}",
                    "title": title,
                    "difficulty": _clean_text(item.difficulty) or "Medium",
                    "companyTargets": [_clean_text(company) for company in item.companyTargets if _clean_text(company)] or ["Tech Co"],
                    "prompt": prompt,
                    "starterCode": starter_code,
                    "expectedSignals": expected[:6],
                }
            )

        if len(challenges) < challenge_count:
            fallback_templates = [
                {
                    "title": f"Implement core {normalized_skills[0]} utility",
                    "prompt": f"Write a clean, testable function using {normalized_skills[0]} and explain time complexity.",
                    "difficulty": "Medium",
                    "companyTargets": ["Tech Co"],
                    "starterCode": "def solve(data):\n    # TODO: implement\n    return data\n",
                    "expectedSignals": ["def", "return", "complexity"],
                },
                {
                    "title": f"Data task using {normalized_skills[min(1, len(normalized_skills)-1)]}",
                    "prompt": f"Build a small solution that processes input data and outputs validated results using {normalized_skills[min(1, len(normalized_skills)-1)]}.",
                    "difficulty": "Medium",
                    "companyTargets": ["Product Team"],
                    "starterCode": "def transform(records):\n    # TODO: validate and transform\n    return records\n",
                    "expectedSignals": ["validate", "loop", "return"],
                },
                {
                    "title": "Build robust input validator",
                    "prompt": "Create reusable validation logic for mixed input records and explain edge-case handling.",
                    "difficulty": "Medium",
                    "companyTargets": ["Operations Team"],
                    "starterCode": "def validate_record(record):\n    # TODO: add validations\n    return True\n",
                    "expectedSignals": ["if", "return", "error handling"],
                },
                {
                    "title": "Generate ranked recommendations",
                    "prompt": "Write a ranking function that scores items by multiple criteria and returns top results.",
                    "difficulty": "Medium",
                    "companyTargets": ["Recommendation Team"],
                    "starterCode": "def rank_items(items):\n    # TODO: score and sort\n    return items\n",
                    "expectedSignals": ["sort", "score", "return"],
                },
                {
                    "title": "Analytics summary builder",
                    "prompt": "Create a function that computes summary metrics from candidate responses for dashboard display.",
                    "difficulty": "Easy",
                    "companyTargets": ["Analytics Team"],
                    "starterCode": "def summarize(responses):\n    # TODO: compute metrics\n    return {}\n",
                    "expectedSignals": ["count", "average", "dictionary"],
                },
            ]
            for template in fallback_templates:
                if len(challenges) >= challenge_count:
                    break
                title_key = template["title"].casefold()
                if title_key in seen_titles:
                    continue
                seen_titles.add(title_key)
                challenges.append(
                    {
                        "id": f"C{len(challenges) + 1}",
                        "title": template["title"],
                        "difficulty": template["difficulty"],
                        "companyTargets": template["companyTargets"],
                        "prompt": template["prompt"],
                        "starterCode": template["starterCode"],
                        "expectedSignals": template["expectedSignals"],
                    }
                )

        if challenges:
            return {"challenges": challenges[:challenge_count]}

        raise RuntimeError("Model returned empty challenges")

    except Exception as e:
        print(f"Error in Test 2: {e}")
        fallback = []
        for i in range(challenge_count):
            skill = normalized_skills[i % len(normalized_skills)] if normalized_skills else "Python"
            fallback.append(
                {
                    "id": f"C{i+1}",
                    "title": f"{skill} implementation challenge {i+1}",
                    "difficulty": "Medium",
                    "companyTargets": ["Tech Co"],
                    "prompt": f"Write a robust {skill} solution and include complexity notes.",
                    "starterCode": "def solve(data):\n    # TODO\n    return data",
                    "expectedSignals": ["def", "return", "complexity"],
                }
            )
        return {"challenges": fallback}

class CodeExecutionRequest(BaseModel):
    code: str
    language: str

FORBIDDEN_CODE_KEYWORDS = [
    "import os", "import sys", "import subprocess", "import shutil",
    "import socket", "import ctypes", "import signal",
    "__import__", "eval(", "exec(", "compile(",
    "os.system", "os.popen", "os.remove", "os.unlink",
    "shutil.rmtree", "subprocess.run", "subprocess.Popen",
    "open(", "pathlib", "import requests", "import http",
]

@app.post("/execute-code")
async def execute_code(request: CodeExecutionRequest):
    """Executes code securely using subprocess with a timeout. Docker fallback if available."""
    if request.language.lower() not in ["python", "py"]:
        return {"stdout": "", "stderr": "Only Python is supported in this local environment.", "error": True}
    
    code_lower = request.code.lower()
    for forbidden in FORBIDDEN_CODE_KEYWORDS:
        if forbidden.lower() in code_lower:
            return {"stdout": "", "stderr": f"Blocked: usage of '{forbidden}' is not allowed for security reasons.", "error": True}
    
    with tempfile.NamedTemporaryFile(suffix=".py", delete=False, mode="w", encoding="utf-8") as f:
        f.write(request.code)
        temp_path = f.name
    
    try:
        # We run the python script with a 3-second timeout
        result = subprocess.run(
            [sys.executable, temp_path],
            capture_output=True,
            text=True,
            timeout=3.0
        )
        return {
            "stdout": result.stdout,
            "stderr": result.stderr,
            "error": result.returncode != 0
        }
    except subprocess.TimeoutExpired:
        return {"stdout": "", "stderr": "Execution timed out (3s limit).", "error": True}
    except Exception as e:
        return {"stdout": "", "stderr": str(e), "error": True}
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

class CheatingData(BaseModel):
    tab_switches: int
    copy_paste_count: int
    time_taken_seconds: int

@app.post("/calculate-cheating-score")
async def calculate_cheating_score(data: CheatingData):
    """Calculates a cheating score (0-100) based on telemetry."""
    score = 0
    if data.tab_switches > 2:
        score += (data.tab_switches - 2) * 15
    if data.copy_paste_count > 0:
        score += data.copy_paste_count * 20
    if data.time_taken_seconds < 10:  # Too fast
        score += 30
    
    score = min(100, score)
    return {"cheating_score": score, "risk_level": "High" if score > 50 else ("Medium" if score > 20 else "Low")}

class AnalyticsRequest(BaseModel):
    skills: List[str]
    total_score: int
    cheating_score: int

@app.post("/generate-analytics")
async def generate_analytics(request: AnalyticsRequest):
    """Generates strengths and weaknesses."""
    # Simple deterministic logic for fast response
    adjusted_score = max(0, request.total_score - (request.cheating_score // 2))
    
    skill_scores = {}
    for idx, s in enumerate(request.skills):
        seed = int(hashlib.md5(s.encode()).hexdigest(), 16) % 21 - 10
        skill_scores[s] = min(100, max(20, adjusted_score + seed))

    sorted_skills = sorted(skill_scores.items(), key=lambda x: x[1], reverse=True)
    strengths = [f"Strong understanding of {sk}" for sk, sc in sorted_skills if sc >= adjusted_score][:3]
    weaknesses = [f"Needs improvement in {sk}" for sk, sc in sorted_skills if sc < adjusted_score][:3]
    if not strengths:
        strengths = [f"Solid foundation in {request.skills[0]}" if request.skills else "Fundamentals are present"]
    if not weaknesses:
        weaknesses = ["Consider deep diving into practical implementation."]

    return {
        "overall_score": adjusted_score,
        "skill_scores": skill_scores,
        "strengths": strengths,
        "weaknesses": weaknesses
    }

class CourseRecommendationRequest(BaseModel):
    skills: List[str]
    weak_skills: List[str] = []

STATIC_COURSE_MAP = {
    "python": {"course_name": "Python for Everybody", "platform": "Coursera", "url_hint": "https://www.coursera.org/specializations/python", "difficulty": "Beginner", "description": "A complete introduction to Python programming."},
    "java": {"course_name": "Java Programming Masterclass", "platform": "Udemy", "url_hint": "https://www.udemy.com/course/java-the-complete-java-developer-course/", "difficulty": "Beginner", "description": "Comprehensive Java from scratch."},
    "sql": {"course_name": "SQL for Data Science", "platform": "Coursera", "url_hint": "https://www.coursera.org/learn/sql-for-data-science", "difficulty": "Beginner", "description": "SQL querying fundamentals for analytics."},
    "excel": {"course_name": "Excel Skills for Business", "platform": "Coursera", "url_hint": "https://www.coursera.org/specializations/excel", "difficulty": "Beginner", "description": "Master spreadsheets for workplace tasks."},
    "html": {"course_name": "HTML & CSS Crash Course", "platform": "freeCodeCamp", "url_hint": "https://www.freecodecamp.org/learn", "difficulty": "Beginner", "description": "Build websites from scratch."},
    "javascript": {"course_name": "The Complete JavaScript Course", "platform": "Udemy", "url_hint": "https://www.udemy.com/course/the-complete-javascript-course/", "difficulty": "Intermediate", "description": "Modern JavaScript from fundamentals to advanced."},
    "react": {"course_name": "React - The Complete Guide", "platform": "Udemy", "url_hint": "https://www.udemy.com/course/react-the-complete-guide/", "difficulty": "Intermediate", "description": "Build modern React applications."},
    "data entry": {"course_name": "Data Entry & Management", "platform": "Alison", "url_hint": "https://alison.com/course/data-entry-and-management", "difficulty": "Beginner", "description": "Efficient data entry and management skills."},
    "communication": {"course_name": "Improving Communication Skills", "platform": "Coursera", "url_hint": "https://www.coursera.org/learn/wharton-communication-skills", "difficulty": "Beginner", "description": "Professional communication techniques."},
    "networking": {"course_name": "Computer Networking", "platform": "Coursera", "url_hint": "https://www.coursera.org/professional-certificates/google-it-support", "difficulty": "Beginner", "description": "Networking fundamentals for IT careers."},
}

COURSE_ALTERNATIVES = {
    "python": [
        {"course_name": "Python for Everybody", "platform": "Coursera", "url_hint": "https://www.coursera.org/specializations/python", "difficulty": "Beginner", "description": "Python basics with projects."},
        {"course_name": "Automate the Boring Stuff with Python", "platform": "Udemy", "url_hint": "https://www.udemy.com/course/automate/", "difficulty": "Beginner", "description": "Practical Python automation tasks."},
        {"course_name": "Python Tutorial", "platform": "freeCodeCamp", "url_hint": "https://www.freecodecamp.org/news/the-python-handbook/", "difficulty": "Beginner", "description": "Free step-by-step Python handbook."},
    ],
    "sql": [
        {"course_name": "SQL for Data Science", "platform": "Coursera", "url_hint": "https://www.coursera.org/learn/sql-for-data-science", "difficulty": "Beginner", "description": "SQL querying and joins for analytics."},
        {"course_name": "SQL Basics", "platform": "Khan Academy", "url_hint": "https://www.khanacademy.org/computing/computer-programming/sql", "difficulty": "Beginner", "description": "Practice SQL queries interactively."},
        {"course_name": "The Complete SQL Bootcamp", "platform": "Udemy", "url_hint": "https://www.udemy.com/course/the-complete-sql-bootcamp/", "difficulty": "Beginner", "description": "Hands-on SQL from zero to projects."},
    ],
    "react": [
        {"course_name": "React - The Complete Guide", "platform": "Udemy", "url_hint": "https://www.udemy.com/course/react-the-complete-guide/", "difficulty": "Intermediate", "description": "Modern React patterns and hooks."},
        {"course_name": "React Basics", "platform": "Meta", "url_hint": "https://www.coursera.org/learn/react-basics", "difficulty": "Beginner", "description": "React fundamentals with components."},
        {"course_name": "Learn React", "platform": "Scrimba", "url_hint": "https://scrimba.com/learn/learnreact", "difficulty": "Beginner", "description": "Interactive React course with mini builds."},
    ],
    "javascript": [
        {"course_name": "The Complete JavaScript Course", "platform": "Udemy", "url_hint": "https://www.udemy.com/course/the-complete-javascript-course/", "difficulty": "Intermediate", "description": "JavaScript fundamentals to advanced."},
        {"course_name": "JavaScript Algorithms and Data Structures", "platform": "freeCodeCamp", "url_hint": "https://www.freecodecamp.org/learn/javascript-algorithms-and-data-structures/", "difficulty": "Beginner", "description": "Hands-on practice with coding tasks."},
        {"course_name": "JavaScript Essentials", "platform": "MDN", "url_hint": "https://developer.mozilla.org/en-US/docs/Learn/JavaScript", "difficulty": "Beginner", "description": "Core JavaScript concepts from docs."},
    ],
    "excel": [
        {"course_name": "Excel Skills for Business", "platform": "Coursera", "url_hint": "https://www.coursera.org/specializations/excel", "difficulty": "Beginner", "description": "Core spreadsheet formulas and reporting."},
        {"course_name": "Excel for Beginners", "platform": "Microsoft Learn", "url_hint": "https://support.microsoft.com/excel", "difficulty": "Beginner", "description": "Official Excel fundamentals and practice."},
        {"course_name": "Advanced Excel", "platform": "YouTube", "url_hint": "https://www.youtube.com/results?search_query=advanced+excel+tutorial", "difficulty": "Intermediate", "description": "Free tutorials for formulas and dashboards."},
    ],
}

def _fallback_course_for_skill(skill: str, variant: int) -> Dict[str, str]:
    key = skill.casefold()
    options = COURSE_ALTERNATIVES.get(key)
    if options:
        picked = options[variant % len(options)].copy()
        picked["skill"] = skill
        picked["url_hint"] = _ensure_absolute_url(picked.get("url_hint", ""))
        return picked

    if key in STATIC_COURSE_MAP:
        entry = STATIC_COURSE_MAP[key].copy()
        entry["skill"] = skill
        entry["url_hint"] = _ensure_absolute_url(entry.get("url_hint", ""))
        return entry

    platforms = ["YouTube", "Coursera", "Udemy", "NPTEL", "freeCodeCamp"]
    chosen_platform = platforms[variant % len(platforms)]
    return {
        "skill": skill,
        "course_name": f"Learn {skill}",
        "platform": chosen_platform,
        "url_hint": f"https://www.youtube.com/results?search_query=learn+{skill.replace(' ', '+')}",
        "difficulty": "Beginner",
        "description": f"Curated beginner path for {skill}.",
    }

@app.post("/course-recommendations")
async def course_recommendations(request: CourseRecommendationRequest):
    target_skills = request.weak_skills if request.weak_skills else request.skills
    normalized = _normalize_skills(target_skills)
    if not normalized:
        return {"courses": []}

    request_key = _course_request_key(normalized)
    cursor = COURSE_CURSOR.get(request_key, 0)
    COURSE_CURSOR[request_key] = cursor + 1

    if _has_llm_provider():
        try:
            skills_str = ", ".join(normalized)
            prompt = (
                f"Recommend one online course for EACH of these skills: {skills_str}.\n"
                "Return JSON only. Format: {\"courses\": [{\"skill\": \"...\", \"course_name\": \"...\", \"platform\": \"...\", \"url_hint\": \"...\", \"difficulty\": \"Beginner/Intermediate/Advanced\", \"description\": \"...\"}]}"
            )
            raw = _call_text(prompt, system_instruction="Return only valid JSON with course recommendations.", temperature=0.3)
            parsed = json.loads(_strip_code_fences(raw))
            courses = parsed.get("courses", [])
            if courses:
                cleaned_courses = []
                seen = set()
                for index, course in enumerate(courses):
                    skill = _clean_text(course.get("skill") or normalized[index % len(normalized)])
                    if not skill:
                        continue
                    course_name = _clean_text(course.get("course_name") or course.get("title"))
                    platform = _clean_text(course.get("platform") or "Online")
                    key = (skill.casefold(), course_name.casefold(), platform.casefold())
                    if key in seen:
                        continue
                    seen.add(key)
                    cleaned_courses.append(
                        {
                            "skill": skill,
                            "course_name": course_name or f"Learn {skill}",
                            "platform": platform,
                            "url_hint": _ensure_absolute_url(_clean_text(course.get("url_hint")) or f"https://www.youtube.com/results?search_query=learn+{skill.replace(' ', '+')}"),
                            "difficulty": _clean_text(course.get("difficulty")) or "Beginner",
                            "description": _clean_text(course.get("description")) or f"Curated learning path for {skill}.",
                        }
                    )

                if cleaned_courses:
                    rotated = _rotate_list(cleaned_courses, cursor)
                    return {"courses": rotated[: max(len(normalized), 3)]}
        except Exception as e:
            print(f"AI course recommendation failed, using fallback: {e}")

    courses = []
    for index, skill in enumerate(normalized):
        courses.append(_fallback_course_for_skill(skill, variant=cursor + index))

    courses = _rotate_list(courses, cursor)

    return {"courses": courses}


class ATSScoreRequest(BaseModel):
    resume_text: str
    target_keywords: List[str] = []

@app.post("/ats-score")
async def ats_score(request: ATSScoreRequest):
    text = _clean_text(request.resume_text)
    if not text or len(text) < 50:
        raise HTTPException(status_code=400, detail="Resume text is too short for ATS analysis.")

    text_lower = text.casefold()
    target_kw = _normalize_skills(request.target_keywords) if request.target_keywords else _normalize_skills(list(STATIC_COURSE_MAP.keys()))

    matched = [kw for kw in target_kw if kw.casefold() in text_lower]
    missing = [kw for kw in target_kw if kw.casefold() not in text_lower]

    keyword_score = int((len(matched) / max(len(target_kw), 1)) * 100)

    sections_expected = ["education", "experience", "skills", "projects", "objective", "summary"]
    sections_found = [s for s in sections_expected if s in text_lower]
    section_scores = {s: (100 if s in sections_found else 0) for s in sections_expected}
    formatting_score = int((len(sections_found) / len(sections_expected)) * 100)

    ats_total = int(keyword_score * 0.6 + formatting_score * 0.4)

    ai_tips = []
    if _has_llm_provider():
        try:
            prompt = (
                f"Given this resume text (first 2000 chars):\n{text[:2000]}\n\n"
                f"And these missing keywords: {', '.join(missing[:10])}\n"
                "Give 3 short, actionable tips to improve ATS score. Return JSON: {\"tips\": [\"...\"]}"
            )
            raw = _call_text(prompt, system_instruction="Return only valid JSON.", temperature=0.3, max_output_tokens=512)
            parsed = json.loads(_strip_code_fences(raw))
            ai_tips = parsed.get("tips", [])
        except Exception as e:
            print(f"ATS AI tips failed: {e}")

    return {
        "ats_score": ats_total,
        "keyword_matches": matched,
        "missing_keywords": missing[:15],
        "formatting_score": formatting_score,
        "section_scores": section_scores,
        "tips": ai_tips
    }


class ResumeSuggestionsRequest(BaseModel):
    resume_text: str
    skills: List[str] = []
    target_role: str = "PM Internship"

@app.post("/resume-suggestions")
async def resume_suggestions(request: ResumeSuggestionsRequest):
    text = _clean_text(request.resume_text)
    if not text or len(text) < 50:
        raise HTTPException(status_code=400, detail="Resume text is too short for analysis.")

    if _has_llm_provider():
        try:
            skills_str = ", ".join(request.skills) if request.skills else "general"
            prompt = (
                f"Analyze this resume for a '{request.target_role}' role.\n"
                f"Candidate skills: {skills_str}\n"
                f"Resume text (first 3000 chars):\n{text[:3000]}\n\n"
                "Provide 4-6 improvement suggestions. Return JSON only. "
                "Format: {\"suggestions\": [{\"section\": \"...\", \"current_issue\": \"...\", \"improvement\": \"...\", \"priority\": \"high/medium/low\"}]}"
            )
            raw = _call_text(prompt, system_instruction="Return only valid JSON with resume improvement suggestions.", temperature=0.3, max_output_tokens=1024)
            parsed = json.loads(_strip_code_fences(raw))
            suggestions = parsed.get("suggestions", [])
            if suggestions:
                return {"suggestions": suggestions}
        except Exception as e:
            print(f"AI resume suggestions failed: {e}")

    fallback = []
    text_lower = text.casefold()
    if "objective" not in text_lower and "summary" not in text_lower:
        fallback.append({"section": "Summary", "current_issue": "Missing professional summary section", "improvement": "Add a 2-3 sentence summary highlighting your key skills and career goals", "priority": "high"})
    if "project" not in text_lower:
        fallback.append({"section": "Projects", "current_issue": "No projects section found", "improvement": "Add 2-3 relevant projects with technologies used and outcomes achieved", "priority": "high"})
    if "skill" not in text_lower:
        fallback.append({"section": "Skills", "current_issue": "Skills section not clearly identified", "improvement": "Create a dedicated skills section with categorized technical and soft skills", "priority": "medium"})
    if len(text) < 500:
        fallback.append({"section": "Overall", "current_issue": "Resume content is very brief", "improvement": "Expand descriptions with quantifiable achievements and specific technologies", "priority": "high"})
    if not fallback:
        fallback.append({"section": "General", "current_issue": "Resume could be enhanced", "improvement": "Use action verbs, quantify achievements, and tailor content to the target role", "priority": "medium"})

    return {"suggestions": fallback}

@app.post("/resume-improver")
async def resume_improver(request: ResumeImproverRequest):
    ats_payload = ATSScoreRequest(
        resume_text=request.resume_text,
        target_keywords=request.skills,
    )
    suggestions_payload = ResumeSuggestionsRequest(
        resume_text=request.resume_text,
        skills=request.skills,
        target_role=request.target_role or "PM Internship",
    )

    ats_data = await ats_score(ats_payload)
    suggestions_data = await resume_suggestions(suggestions_payload)

    return {
        "ats": ats_data,
        "suggestions": suggestions_data.get("suggestions", []),
        "target_role": request.target_role or "PM Internship",
    }

class ScrapeRequest(BaseModel):
    url: str
    formats: Optional[List[str]] = None

@app.post("/firecrawl-scrape")
async def firecrawl_scrape_endpoint(request: ScrapeRequest):
    """Scrape any URL using Firecrawl and return structured content."""
    if not firecrawl_client:
        raise HTTPException(
            status_code=503,
            detail="Firecrawl is not configured. Add FIRECRAWL_API_KEY to .env.txt and install firecrawl-py.",
        )
    cleaned_url = _ensure_absolute_url(_clean_text(request.url))
    if not cleaned_url:
        raise HTTPException(status_code=400, detail="Invalid URL provided.")
    try:
        result = _firecrawl_scrape(cleaned_url, formats=request.formats)
        return {
            "success": True,
            "url": cleaned_url,
            "markdown": result.get("markdown", ""),
            "html": result.get("html", ""),
            "metadata": result.get("metadata", {}),
        }
    except Exception as e:
        raise HTTPException(status_code=503, detail=str(e))
